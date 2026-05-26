"""One-off helper: minimal blank audit form for TS pipeline smoke tests."""
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "test-data" / "blank-form.docx"

LABELS = [
    "Facility Name",
    "Activity Code",
    "NDT Personnel",
    "Procedure Reference",
    "Quality Records",
]


def main() -> None:
    doc = Document()
    doc.add_heading("Audit Form — Mock Template", level=1)
    table = doc.add_table(rows=len(LABELS), cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(LABELS):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = ""
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
