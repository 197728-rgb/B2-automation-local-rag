from __future__ import annotations
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from shared_core.file_utils import normalize_label, normalize_ws, ensure_dir, save_json
from shared_core.mapping import FORM_TEMPLATES, FIELD_ALIASES
from shared_core.template_registry import default_registry_key

def build_alias_to_field() -> dict:
    d = {}
    for field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            d[normalize_label(alias)] = field
    return d

ALIAS_TO_FIELD = build_alias_to_field()

def field_for_label(label: str) -> str | None:
    norm = normalize_label(label)
    if not norm:
        return None
    if norm in ALIAS_TO_FIELD:
        return ALIAS_TO_FIELD[norm]
    for alias, field in ALIAS_TO_FIELD.items():
        if alias in norm or norm in alias:
            return field
    return None

def render_value(field: str, data: dict) -> str:
    if field == "tank_spec" and data.get("tank_spec") and data.get("stencil_spec"):
        if data["stencil_spec"] not in data["tank_spec"]:
            return f"{data['tank_spec']} / {data['stencil_spec']}"
    if field == "aar_form_no":
        if data.get("aar_form") and data.get("aar_no"):
            return f"{data['aar_form']} / {data['aar_no']}"
    return str(data.get(field, "") or "")

def unique_cells(cells: list[_Cell]) -> list[_Cell]:
    seen = set()
    out = []
    for cell in cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            out.append(cell)
    return out

def set_cell_text(cell: _Cell, value: str) -> None:
    cell.text = value

class OptimizedFillerEngine:
    def __init__(self, forms_dir: str | Path):
        self.forms_dir = Path(forms_dir)

    def _fill_blank_following_rows(self, doc: Document, data: dict) -> set[str]:
        filled = set()
        for table in doc.tables:
            rows = table.rows
            for r in range(len(rows) - 1):
                header_cells = rows[r].cells
                value_cells = rows[r + 1].cells
                # Build contiguous header groups
                groups = []
                start = 0
                prev = normalize_label(header_cells[0].text) if header_cells else ""
                for i in range(1, len(header_cells) + 1):
                    curr = normalize_label(header_cells[i].text) if i < len(header_cells) else None
                    if i == len(header_cells) or curr != prev:
                        groups.append((start, i, prev))
                        start = i
                        prev = curr
                # Fill next-row blank segments
                for start, end, header_label in groups:
                    field = field_for_label(header_label)
                    if not field:
                        continue
                    value = render_value(field, data)
                    if not value:
                        continue
                    segment = unique_cells(value_cells[start:end])
                    header_segment = unique_cells(header_cells[start:end])
                    if not segment:
                        continue
                    segment_text = normalize_ws(" ".join(c.text for c in segment))
                    # only fill truly blank or placeholder-ish rows
                    if segment_text:
                        continue
                    set_cell_text(segment[0], value)
                    filled.add(field)
        return filled

    def _fill_inline_label_value_rows(self, doc: Document, data: dict) -> set[str]:
        filled = set()
        for table in doc.tables:
            for row in table.rows:
                cells = unique_cells(row.cells)
                labels = [normalize_label(c.text) for c in cells]
                for i, label in enumerate(labels):
                    field = field_for_label(label)
                    if not field:
                        continue
                    value = render_value(field, data)
                    if not value:
                        continue
                    # Find nearest blank cell to the right
                    for j in range(i + 1, len(cells)):
                        if normalize_ws(cells[j].text):
                            continue
                        set_cell_text(cells[j], value)
                        filled.add(field)
                        break
        return filled

    def _fill_marking_tables(self, doc: Document, data: dict) -> set[str]:
        filled = set()
        label_value_fields = {
            "manufacturer’s model number": "manufacturer",
            "manufacturer’s design or type number": "model_number",
            "unique serial number": "serial_number",
            "pressure rating": "pressure_rating",
            "set pressure": "pressure_rating",
            "official flow capacity": "official_flow_capacity",
            "flow rating pressure": "flow_rating_pressure",
            "specific gravity": "specific_gravity",
            "station stencil": "station_stencil",
        }
        for table in doc.tables:
            for row in table.rows:
                cells = unique_cells(row.cells)
                if len(cells) < 2:
                    continue
                left = normalize_label(cells[0].text)
                for key_alias, field in label_value_fields.items():
                    if key_alias in left and data.get(field):
                        # Often observed entry begins around the second or third unique cell
                        target = cells[1] if len(cells) > 1 else cells[0]
                        if not normalize_ws(target.text):
                            set_cell_text(target, str(data[field]))
                        filled.add(field)
        return filled

    def fill_form(self, form_code: str, extracted_data: dict, output_path: str | Path) -> dict:
        dk = default_registry_key()
        template_name = FORM_TEMPLATES.get(form_code) or FORM_TEMPLATES.get(dk, next(iter(FORM_TEMPLATES.values())))
        template_path = self.forms_dir / template_name
        doc = Document(str(template_path))

        filled = set()
        filled |= self._fill_blank_following_rows(doc, extracted_data)
        filled |= self._fill_inline_label_value_rows(doc, extracted_data)
        filled |= self._fill_marking_tables(doc, extracted_data)

        output_path = Path(output_path)
        ensure_dir(output_path.parent)
        doc.save(str(output_path))

        unresolved = sorted(k for k, v in extracted_data.items() if v and k not in filled)
        audit = {
            "template": template_name,
            "form_code": form_code,
            "filled_fields": sorted(filled),
            "unresolved_fields": unresolved,
            "extracted_data": extracted_data,
        }
        save_json(output_path.with_suffix(".json"), audit)
        return {"output_path": str(output_path), "audit_path": str(output_path.with_suffix('.json')), "filled_fields": sorted(filled), "unresolved_fields": unresolved}
