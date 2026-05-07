from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from shared_core.mapping import FORM_TEMPLATES
from shared_core.template_registry import default_registry_key


def validate_b2_output(
    *,
    forms_dir: Path,
    form_code: str,
    output_path: Path,
    filled_fields: list[str],
) -> dict[str, Any]:
    """
    Bible-aligned minimal validation: template vs output table count, file exists,
    non-zero size. Extend with per-cell run counts when writer logs coordinates.
    """
    dk = default_registry_key()
    template_name = FORM_TEMPLATES.get(form_code) or FORM_TEMPLATES.get(dk, next(iter(FORM_TEMPLATES.values())))
    template_path = forms_dir / template_name
    report: dict[str, Any] = {
        "output_exists": output_path.is_file(),
        "output_bytes": output_path.stat().st_size if output_path.is_file() else 0,
        "template_exists": template_path.is_file(),
        "filled_field_count": len(filled_fields),
    }
    if not report["output_exists"] or not report["template_exists"]:
        report["table_count_ok"] = None
        return report

    try:
        t_doc = Document(str(template_path))
        o_doc = Document(str(output_path))
        report["template_tables"] = len(t_doc.tables)
        report["output_tables"] = len(o_doc.tables)
        report["table_count_preserved"] = len(t_doc.tables) == len(o_doc.tables)
        report["table_count_ok"] = report["table_count_preserved"]
    except Exception as exc:
        report["table_count_ok"] = False
        report["validation_error"] = str(exc)
    return report
