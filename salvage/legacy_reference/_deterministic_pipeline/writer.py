"""
writer.py — Run-safe in-place DOCX updates with validation.

Opens DOCX, locates exact targets via resolved bridge_map entries,
updates cells safely at the run level, clones rows when allowed,
emits a validation record for every write.

Does NOT reinterpret evidence or guess values.
"""
from __future__ import annotations

import copy
import json
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import docx
from docx.oxml.ns import qn

from bridge_map import resolve_targets, _find_table, CellTarget


def _count_runs(cell) -> int:
    return sum(len(p.runs) for p in cell.paragraphs)


def _count_paragraphs(cell) -> int:
    return len(cell.paragraphs)


def _write_run_safe(cell, new_value: str, mode: str) -> dict:
    """
    Write new_value into cell using the safest available method.
    Returns integrity info for validation.
    """
    before_pcount = _count_paragraphs(cell)
    before_rcount = _count_runs(cell)
    old_text = cell.text.strip()

    if mode == "replace_run_text":
        written = False
        for para in cell.paragraphs:
            if para.runs:
                para.runs[0].text = new_value
                for extra_run in para.runs[1:]:
                    extra_run.text = ""
                written = True
                break
        if not written and cell.paragraphs:
            cell.paragraphs[0].add_run(new_value)
            written = True

    elif mode == "replace_paragraph_text":
        if cell.paragraphs and cell.paragraphs[0].runs:
            for r in cell.paragraphs[0].runs[1:]:
                r.text = ""
            cell.paragraphs[0].runs[0].text = new_value
        elif cell.paragraphs:
            cell.paragraphs[0].add_run(new_value)

    elif mode == "replace_cell_text":
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = new_value
            for r in cell.paragraphs[0].runs[1:]:
                r.text = ""
            for p in cell.paragraphs[1:]:
                for r in p.runs:
                    r.text = ""
        else:
            cell.text = new_value

    after_pcount = _count_paragraphs(cell)
    after_rcount = _count_runs(cell)

    return {
        "old_text": old_text,
        "new_text": new_value,
        "paragraphs_before": before_pcount,
        "paragraphs_after": after_pcount,
        "runs_before": before_rcount,
        "runs_after": after_rcount,
        "paragraph_delta": after_pcount - before_pcount,
        "run_delta": after_rcount - before_rcount,
    }


def _clone_row(table, source_row_index: int, insert_before_index: int | None = None):
    """Clone a row's XML and insert it before a target row, preserving styling."""
    source_tr = table.rows[source_row_index]._tr
    new_tr = copy.deepcopy(source_tr)
    for tc in new_tr.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = ""

    tbl = table._tbl
    if insert_before_index is not None and insert_before_index < len(table.rows):
        ref_tr = table.rows[insert_before_index]._tr
        ref_tr.addprevious(new_tr)
    else:
        tbl.append(new_tr)
    return new_tr


def apply_writes(
    docx_path: str | Path,
    mapping: list[dict],
    output_path: str | Path | None = None,
) -> dict[str, Any]:
    """
    Apply bridge_map entries to a DOCX file.

    Parameters
    ----------
    docx_path : path to the source DOCX (will not be modified)
    mapping : list of dicts, each with canonical_key, display_value, and target
    output_path : where to save the result (default: <stem>_updated.docx)

    Returns
    -------
    dict with write_log, summary counts, and output path.
    """
    docx_path = Path(docx_path)
    if output_path is None:
        output_path = docx_path.parent / f"{docx_path.stem}_updated.docx"
    output_path = Path(output_path)

    doc = docx.Document(str(docx_path))

    table_count_before = len(doc.tables)
    row_counts_before = {i: len(t.rows) for i, t in enumerate(doc.tables)}

    write_mapping = [e for e in mapping if not e.get("insert_only")]
    resolved = resolve_targets(doc, write_mapping)

    write_log: list[dict] = []
    written = 0
    skipped = 0
    failed = 0

    for entry in resolved:
        log_entry = {
            "canonical_key": entry.get("canonical_key", "?"),
            "display_value": entry.get("display_value", "?"),
            "source_file": entry.get("source_file", ""),
            "source_page": entry.get("source_page", ""),
            "status": entry["status"],
        }

        if entry["status"] != "resolved":
            log_entry["reason"] = entry["status"]
            skipped += 1
            write_log.append(log_entry)
            continue

        cell = entry["cell"]
        new_value = entry["display_value"]
        mode = entry.get("write_mode", "replace_run_text")

        integrity = _write_run_safe(cell, new_value, mode)

        log_entry.update({
            "table_index": entry["table_index"],
            "row_index": entry["row_index"],
            "col_index": entry["col_index"],
            "write_mode": mode,
            **integrity,
        })
        log_entry["status"] = "written"
        written += 1
        write_log.append(log_entry)

    table_count_after = len(doc.tables)
    # Process row insertions (entries with insert_row_data)
    # Group by table anchor so we can handle index shifting within each table
    insert_entries = [e for e in mapping if e.get("insert_only") and e.get("insert_row_data")]
    insert_offset = {}  # track how many rows we've inserted per table

    for entry in insert_entries:
        ird = entry["insert_row_data"]
        anchor = ird.get("table_anchor", "")
        found = _find_table(doc, anchor, ird.get("table_index_hint"))
        if not found:
            write_log.append({
                "canonical_key": entry.get("canonical_key", "?"),
                "status": "insert_table_not_found",
            })
            skipped += 1
            continue
        ti, table = found
        offset = insert_offset.get(ti, 0)

        clone_from = ird.get("clone_from_row", len(table.rows) - 2)
        insert_before = ird.get("insert_before_row")
        if insert_before is not None:
            insert_before += offset

        new_tr = _clone_row(table, clone_from, insert_before)

        from docx.oxml.ns import qn as _qn
        tcs = new_tr.findall(_qn("w:tc"))
        col_values = ird.get("col_values", {})
        for col_str, val in col_values.items():
            ci = int(col_str)
            if ci < len(tcs):
                tc = tcs[ci]
                paras = tc.findall(_qn("w:p"))
                if paras:
                    runs = paras[0].findall(_qn("w:r"))
                    if runs:
                        ts = runs[0].findall(_qn("w:t"))
                        if ts:
                            ts[0].text = val
                        else:
                            from lxml import etree
                            t_el = etree.SubElement(runs[0], _qn("w:t"))
                            t_el.text = val
                    else:
                        from lxml import etree
                        r_el = etree.SubElement(paras[0], _qn("w:r"))
                        t_el = etree.SubElement(r_el, _qn("w:t"))
                        t_el.text = val

        insert_offset[ti] = offset + 1
        write_log.append({
            "canonical_key": entry.get("canonical_key", "?"),
            "status": "row_inserted",
            "table_index": ti,
            "col_values": col_values,
        })
        written += 1

    doc.save(str(output_path))

    row_counts_after = {i: len(t.rows) for i, t in enumerate(doc.tables)}

    return {
        "output_path": str(output_path),
        "written_at": datetime.now(timezone.utc).isoformat(),
        "source_docx": str(docx_path),
        "summary": {
            "total_entries": len(mapping),
            "resolved": written + skipped,
            "written": written,
            "skipped": skipped,
            "failed": failed,
        },
        "structural_integrity": {
            "table_count_before": table_count_before,
            "table_count_after": table_count_after,
            "table_count_changed": table_count_before != table_count_after,
            "row_count_changes": {
                str(k): {"before": row_counts_before.get(k, 0), "after": v}
                for k, v in row_counts_after.items()
                if row_counts_before.get(k, 0) != v
            },
        },
        "write_log": write_log,
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python writer.py <source.docx> <mapping.json> [output.docx]")
        raise SystemExit(1)

    src = Path(sys.argv[1])
    with open(sys.argv[2], encoding="utf-8") as f:
        mapping = json.load(f)
    out = Path(sys.argv[3]) if len(sys.argv) > 3 else None

    result = apply_writes(src, mapping, out)
    log_path = Path(result["output_path"]).with_suffix(".write_log.json")
    log_path.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")

    s = result["summary"]
    print(f"Written: {s['written']}, Skipped: {s['skipped']}, Failed: {s['failed']}")
    print(f"Output: {result['output_path']}")
    print(f"Log: {log_path}")
