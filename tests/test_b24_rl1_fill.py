"""B24_RL1: manifest-driven partial fill on real template."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from b2_automation.b24_rl1_filler import fill_b24_rl1_partial, load_manifest


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _all_doc_text(doc: Document) -> str:
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _cell_text_for_field(doc: Document, manifest: dict, field_id: str) -> str:
    for spec in manifest["cells"]:
        if spec["field_id"] == field_id:
            t, r, c = int(spec["table_index"]), int(spec["row"]), int(spec["col"])
            return doc.tables[t].rows[r].cells[c].text
    raise KeyError(field_id)


@pytest.fixture
def b24_paths():
    root = _repo_root()
    template = root / "templates" / "B24_RL1.docx"
    manifest = root / "schemas" / "templates" / "B24_RL1.json"
    return root, template, manifest


def test_b24_rl1_partial_fill_writes_expected_content(b24_paths, tmp_path: Path) -> None:
    root, template, manifest_path = b24_paths
    if not template.is_file():
        pytest.skip(f"missing template: {template}")
    manifest = load_manifest(manifest_path)
    fields = {
        "tco_name": "Demo TCO LLC",
        "tco_permission_date": "2026-01-15",
        "tco_written_instructions": "Written OK for RL1 demo scope.",
        "pitp_document_name": "PITP-2025-001",
        "pitp_id": "PITP-DEMO-42",
        "pitp_approved_by": "A. Reviewer",
        "pitp_date_approved": "2026-01-10",
        "pitp_rev": "B",
        "car_mark": "DOTX 123456",
        "tank_design_spec": "Stencil DEMO-01",
        "car_type": "DOT-117J100W",
        "aar_form_4_2_number": "AAR-42-DEMO",
        "four_two_drawing_number": "DWG-DEMO-1001",
        "four_two_drawing_revision": "Rev 2",
        "test_plate_tank_material": "TC-128 Grade B",
        "test_plate_tank_mtr": "MTR-DEMO-001",
        "attachment_material": "A516-70 pad",
        "evidence_notes": (
            "Source: cover sheet p1 (TCO), stencil photo p3 (car mark). "
            "Confidence: TCO 0.91, car mark 0.88, PITP ID 0.85."
        ),
    }
    out = tmp_path / "B24_RL1_filled.docx"
    fill_b24_rl1_partial(template, manifest, fields, out)
    assert out.is_file()
    assert out.stat().st_size > 10_000
    doc = Document(out)
    blob = _all_doc_text(doc)
    assert "Demo TCO LLC" in blob
    assert "DOTX 123456" in blob
    assert "PITP-2025-001" in blob
    assert "Written OK for RL1 demo scope." in blob
    assert "PITP-DEMO-42" in blob
    assert "A. Reviewer" in blob
    assert "2026-01-10" in blob
    assert "Stencil DEMO-01" in blob
    assert "DWG-DEMO-1001" in blob
    assert "MTR-DEMO-001" in blob
    assert "A516-70 pad" in blob
    assert "confidence" in blob.lower()
    assert "source" in blob.lower()


def test_partial_fill_omitted_keys_preserve_template_cells(b24_paths, tmp_path: Path) -> None:
    """Regression: do not write empty strings into cells for missing field_ids."""
    _root, template, manifest_path = b24_paths
    if not template.is_file():
        pytest.skip(f"missing template: {template}")
    manifest = load_manifest(manifest_path)
    baseline = Document(template)
    before_pitp = _cell_text_for_field(baseline, manifest, "pitp_id")
    before_car_type = _cell_text_for_field(baseline, manifest, "car_type")
    out = tmp_path / "B24_RL1_minimal_partial.docx"
    fill_b24_rl1_partial(
        template,
        manifest,
        {"tco_name": "MINIMAL_PARTIAL_TCO_ONLY"},
        out,
    )
    filled = Document(out)
    assert _cell_text_for_field(filled, manifest, "tco_name") == "MINIMAL_PARTIAL_TCO_ONLY"
    assert _cell_text_for_field(filled, manifest, "pitp_id") == before_pitp
    assert _cell_text_for_field(filled, manifest, "car_type") == before_car_type


def test_partial_fill_empty_string_key_clears_cell(b24_paths, tmp_path: Path) -> None:
    """Explicit ``""`` is a real value: it clears the cell; omitted keys still preserve others."""
    _, template, manifest_path = b24_paths
    if not template.is_file():
        pytest.skip(f"missing template: {template}")
    manifest = load_manifest(manifest_path)
    seeded = tmp_path / "B24_RL1_seeded.docx"
    fill_b24_rl1_partial(
        template,
        manifest,
        {"tco_name": "SEED_FOR_CLEAR_TEST"},
        seeded,
    )
    mid_doc = Document(seeded)
    assert _cell_text_for_field(mid_doc, manifest, "tco_name") == "SEED_FOR_CLEAR_TEST"
    before_pitp = _cell_text_for_field(mid_doc, manifest, "pitp_id")
    cleared = tmp_path / "B24_RL1_cleared.docx"
    fill_b24_rl1_partial(
        seeded,
        manifest,
        {"tco_name": ""},
        cleared,
    )
    out_doc = Document(cleared)
    assert _cell_text_for_field(out_doc, manifest, "tco_name") == ""
    assert _cell_text_for_field(out_doc, manifest, "pitp_id") == before_pitp
