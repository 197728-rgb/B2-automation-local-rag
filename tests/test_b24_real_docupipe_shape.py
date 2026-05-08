"""Acceptance: realistic DocuPipe JSON -> normalizer -> B24_RL1 cells -> DOCX with value + confidence + source."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from b2_automation.b24_normalizer import normalize_docupipe_payload_for_b24_rl1
from b2_automation.b24_pipeline import run_b24_rl1_from_docupipe
from b2_automation.paths import B24_SHARED_TEMPLATE_DOCX

pytestmark = pytest.mark.legacy_rl1


def _repo() -> Path:
    return Path(__file__).resolve().parents[1]


def _fixture() -> Path:
    return _repo() / "samples" / "docupipe" / "realistic_b24_response.json"


def _all_doc_text(doc: Document) -> str:
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_normalizer_maps_realistic_docupipe_keys() -> None:
    raw = __import__("json").loads(_fixture().read_text(encoding="utf-8"))
    fields = normalize_docupipe_payload_for_b24_rl1(raw)
    assert fields["tco_name"] == "Midwest Tank Rail Inc"
    assert fields["tco_permission_date"] == "2025-11-20"
    assert fields["pitp_document_name"] == "PITP MC-2024-77 Rev C"
    assert fields["car_mark"] == "UTLX 556891"
    assert fields["tco_written_instructions"].startswith("RL1 repair")
    assert fields["pitp_id"] == "MC-2024-77"
    assert fields["four_two_drawing_number"] == "DWG-UTLX-4419"
    assert fields["test_plate_tank_mtr"] == "MTR-99821-A"
    assert "confidence=0.94" in fields["evidence_notes"]
    assert "source: page_index=0" in fields["evidence_notes"]
    assert "PITPIdentifier" in fields["evidence_notes"]
    assert "confidence=0.90" in fields["evidence_notes"]


def test_pipeline_writes_docx_with_evidence(tmp_path: Path) -> None:
    root = _repo()
    tpl = root / "templates" / B24_SHARED_TEMPLATE_DOCX
    if not tpl.is_file():
        pytest.skip(f"missing template: {tpl}")
    out = tmp_path / "b24_from_realistic_docupipe.docx"
    run_b24_rl1_from_docupipe(_fixture(), root, out)
    assert out.is_file() and out.stat().st_size > 10_000
    blob = _all_doc_text(Document(out))
    assert "Midwest Tank Rail Inc" in blob
    assert "UTLX 556891" in blob
    assert "PITP MC-2024-77 Rev C" in blob
    assert "MC-2024-77" in blob
    assert "J. Ortega" in blob
    assert "DWG-UTLX-4419" in blob
    assert "MTR-99821-A" in blob
    assert "TC-128 Grade B normalized plate" in blob
    assert "confidence" in blob.lower()
    assert "source" in blob.lower()
