"""SPEC-1 autonomous pipeline tests."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from b2_automation.analyst_agent import analyze_blank_form
from b2_automation.autonomous_pipeline import run_autonomous_template
from b2_automation.investigator_agent import gather_evidence
from b2_automation.schema_catalog import load_available_schemas
from b2_automation.validation_gate import validate_answer
from b2_automation.writer_agent import synthesize_human_response


def _minimal_audit_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Audit Form Test")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Facility Name"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "Date"
    table.rows[1].cells[1].text = ""
    doc.save(path)


def test_schema_catalog_loads_b24_paths():
    root = Path(__file__).resolve().parents[1]
    catalog = load_available_schemas(root, ("B24_RL2",))
    assert catalog
    assert any("demonstration" in p for entry in catalog for p in entry.get("paths", []))


def test_analyze_blank_form_deterministic(tmp_path: Path):
    docx = tmp_path / "B24_RL2.docx"
    _minimal_audit_docx(docx)
    root = Path(__file__).resolve().parents[1]
    field_map = analyze_blank_form(docx, root=root, form_id="B24_RL2", use_llm=False)
    assert field_map.version == "machine_field_map.v1"
    assert field_map.fields
    assert field_map.summary.detected_field_count >= 1
    req = field_map.fields[0]
    assert req.form_location.table_index is not None


def test_gather_evidence_empty_inbox(tmp_path: Path):
    docx = tmp_path / "form.docx"
    _minimal_audit_docx(docx)
    root = Path(__file__).resolve().parents[1]
    field_map = analyze_blank_form(docx, root=root, use_llm=False)
    req = field_map.fields[0]
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    bundle = gather_evidence(req, inbox, form_id="B24_RL2")
    assert bundle.requirement_id == req.id
    assert bundle.gaps


def test_writer_never_review_required(tmp_path: Path):
    docx = tmp_path / "form.docx"
    _minimal_audit_docx(docx)
    root = Path(__file__).resolve().parents[1]
    field_map = analyze_blank_form(docx, root=root, use_llm=False)
    req = field_map.fields[0]
    from b2_automation.autonomous_contracts import EvidenceBundle

    bundle = EvidenceBundle(requirement_id=req.id, gaps=["no files"])
    answer = synthesize_human_response(req, bundle)
    validated = validate_answer(req, bundle, answer, template_path=str(docx))
    assert "REVIEW_REQUIRED" not in validated.text.upper()
    assert validated.fallback_applied


def test_autonomous_template_e2e(tmp_path: Path):
    root = Path(__file__).resolve().parents[1]
    docx = tmp_path / "B24_RL2.docx"
    _minimal_audit_docx(docx)
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "evidence.txt").write_text(
        "Facility Name: Midwest Tank Rail Inc. Date: 05/18/2026 demonstration station Taller México",
        encoding="utf-8",
    )
    manifest_src = root / "schemas" / "templates" / "B24_RL2.json"
    if manifest_src.is_file():
        (tmp_path / "schemas" / "templates").mkdir(parents=True, exist_ok=True)
        import shutil

        shutil.copy(manifest_src, tmp_path / "schemas" / "templates" / "B24_RL2.json")

    out = tmp_path / "out"
    outcome = run_autonomous_template(
        docx,
        source_folder=inbox,
        output_dir=out,
        root=root,
        form_id="B24_RL2",
        use_llm_analyst=False,
        persist_sqlite=False,
    )
    assert outcome.completed_docx
    assert Path(outcome.completed_docx).is_file()
    manifest = out / "run_manifest.json"
    audit_map = list((out / "audit-trail").glob("*_machine_field_map.v1.json"))
    assert audit_map or (out / "audit-trail").exists()
    assert "REVIEW_REQUIRED" not in Path(outcome.completed_docx).read_bytes().decode("utf-8", errors="ignore")


def test_machine_field_map_schema_file_exists():
    root = Path(__file__).resolve().parents[1]
    schema_path = root / "schemas" / "contracts" / "machineFieldMap.v1.schema.json"
    assert schema_path.is_file()
    data = json.loads(schema_path.read_text(encoding="utf-8"))
    assert data["$id"] == "machine_field_map.v1"
