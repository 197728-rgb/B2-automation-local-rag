"""Smoke tests to verify the development environment is working."""

import importlib


def test_python_docx_import():
    mod = importlib.import_module("docx")
    assert hasattr(mod, "Document")


def test_dotenv_import():
    mod = importlib.import_module("dotenv")
    assert hasattr(mod, "load_dotenv")


def test_requests_import():
    import requests
    assert hasattr(requests, "get")


def test_b2_package_import():
    import b2_automation

    assert b2_automation.__version__


def test_create_docx_document():
    from docx import Document

    doc = Document()
    doc.add_heading("B2 Objective Evidence", level=1)
    doc.add_paragraph("Smoke test paragraph.")
    assert len(doc.paragraphs) == 2
