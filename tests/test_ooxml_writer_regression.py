"""Regression coverage for raw OOXML DOCX production writes."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from b2_automation.ooxml_writer import count_docx_structure, patch_docx_cells
from b2_automation.paths import B24_SHARED_TEMPLATE_DOCX


def _repo() -> Path:
    return Path(__file__).resolve().parents[1]


def _load_manifest(root: Path) -> dict:
    path = root / "schemas" / "templates" / "B24_RL2.json"
    return json.loads(path.read_text(encoding="utf-8"))


def _doc_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_ooxml_patch_preserves_table_cell_structure(tmp_path: Path) -> None:
    root = _repo()
    template = root / "templates" / B24_SHARED_TEMPLATE_DOCX
    if not template.is_file():
        pytest.skip(f"missing template: {template}")
    manifest = _load_manifest(root)
    out = tmp_path / "ooxml_filled.docx"
    guard_path = tmp_path / "structure_guard_report.json"

    outcome = patch_docx_cells(
        template,
        manifest,
        {"facility_name": "OOXML Facility", "car_mark": "UTLX 556891"},
        out,
        structure_guard_report_path=guard_path,
    )

    assert outcome.structure_guard_passed is True
    assert guard_path.is_file()
    guard = json.loads(guard_path.read_text(encoding="utf-8"))
    assert guard["pass"] is True
    before = count_docx_structure(template)
    after = count_docx_structure(out)
    for key in ("tables", "rows", "cells", "gridSpan", "vMerge"):
        assert after[key] == before[key]
    assert after["text_nodes"] - before["text_nodes"] == guard["text_nodes_delta_expected"]
    blob = _doc_text(out)
    assert "OOXML Facility" in blob
    assert "UTLX 556891" in blob


def test_ooxml_fixture_like_values_round_trip(tmp_path: Path) -> None:
    """Patch DOCX using RL2 manifest field ids (same template as production B24_RL2)."""
    root = _repo()
    template = root / "templates" / B24_SHARED_TEMPLATE_DOCX
    if not template.is_file():
        pytest.skip(f"missing template: {template}")
    manifest = _load_manifest(root)
    fields = {
        "facility_name": "Midwest Tank Rail Inc",
        "pitp_id": "MC-2024-77",
        "car_mark": "UTLX 556891",
        "car_type": "DOT-117J100W",
        "tco_written_instructions": "Repair authorized per email 2025-11-18.",
    }
    out = tmp_path / "normalized_values.docx"
    outcome = patch_docx_cells(root / "templates" / B24_SHARED_TEMPLATE_DOCX, manifest, fields, out)

    assert outcome.structure_guard_passed is True
    blob = _doc_text(out)
    assert "Midwest Tank Rail Inc" in blob
    assert "UTLX 556891" in blob


def test_exact_approval_map_coordinate_mismatch_fails_guard(tmp_path: Path) -> None:
    root = _repo()
    template = root / "templates" / B24_SHARED_TEMPLATE_DOCX
    if not template.is_file():
        pytest.skip(f"missing template: {template}")
    manifest = _load_manifest(root)
    approval = {
        "fields": {
            "facility_name": {
                "field_id": "facility_name",
                "table_index": 99,
                "row": 4,
                "col": 0,
                "label": "wrong table index",
            }
        }
    }
    out = tmp_path / "denied.docx"
    guard_path = tmp_path / "structure_guard_report.json"
    outcome = patch_docx_cells(
        template,
        manifest,
        {"facility_name": "Should Not Hand Off"},
        out,
        approval_map=approval,
        structure_guard_report_path=guard_path,
    )
    assert outcome.structure_guard_passed is False
    assert not out.is_file()
    guard = json.loads(guard_path.read_text(encoding="utf-8"))
    assert guard["pass"] is False
