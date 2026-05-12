"""Approval map loading, structure guard handoff, and no-fallback regression tests (Stage 6+7)."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from b2_automation.approval_maps import load_exact_approval_bundle, load_exact_approval_bundle_checked
from b2_automation.inbox_pipeline import run_inbox_pipeline
from b2_automation.paths import B24_SHARED_TEMPLATE_DOCX

FIRST_CLASS_FORMS = ("B24_RL2", "B81", "B89", "B90", "Cover_Page")


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


# ---------------------------------------------------------------------------
# Per-form exact map loading
# ---------------------------------------------------------------------------

class TestPerFormMapLoading:
    @pytest.mark.parametrize("form_id", FIRST_CLASS_FORMS)
    def test_map_loads_for_first_class_form(self, form_id: str) -> None:
        root = _repo_root()
        bundle = load_exact_approval_bundle(root, form_id)
        assert bundle is not None, f"No approval bundle for {form_id}"
        assert bundle.map_path.name == f"{form_id}.json"
        assert bundle.approval_map["form_id"] == form_id
        assert bundle.approval_map["form_version"] == "2026"
        fields = bundle.approval_map.get("fields") or {}
        assert len(fields) >= 2, f"{form_id}: map has too few fields ({len(fields)})"

    @pytest.mark.parametrize("form_id", FIRST_CLASS_FORMS)
    def test_map_fields_have_exact_coordinates(self, form_id: str) -> None:
        root = _repo_root()
        bundle = load_exact_approval_bundle(root, form_id)
        assert bundle is not None
        for fid, spec in bundle.approval_map["fields"].items():
            assert "table_index" in spec, f"{form_id}/{fid}: missing table_index"
            assert "row" in spec, f"{form_id}/{fid}: missing row"
            assert "col" in spec, f"{form_id}/{fid}: missing col"
            assert spec["field_id"] == fid

    @pytest.mark.parametrize("form_id", FIRST_CLASS_FORMS)
    def test_manifest_matches_map_fields(self, form_id: str) -> None:
        root = _repo_root()
        bundle = load_exact_approval_bundle(root, form_id)
        assert bundle is not None
        map_fids = set(bundle.approval_map["fields"].keys())
        manifest_fids = {str(c["field_id"]) for c in bundle.manifest.get("cells", [])}
        assert map_fids == manifest_fids, f"{form_id}: map/manifest field_id mismatch: {map_fids ^ manifest_fids}"

    @pytest.mark.parametrize("form_id", FIRST_CLASS_FORMS)
    def test_no_duplicate_coordinates_in_map(self, form_id: str) -> None:
        root = _repo_root()
        result = load_exact_approval_bundle_checked(root, form_id)
        dup_errors = [e for e in result.errors if "duplicate coordinate" in e]
        assert not dup_errors, f"{form_id}: {dup_errors}"

    @pytest.mark.parametrize("form_id", FIRST_CLASS_FORMS)
    def test_no_duplicate_merged_physical_cells_in_map(self, form_id: str) -> None:
        root = _repo_root()
        result = load_exact_approval_bundle_checked(root, form_id)
        merged_errors = [e for e in result.errors if "duplicate physical cell" in e]
        assert not merged_errors, f"{form_id}: {merged_errors}"


# ---------------------------------------------------------------------------
# Version / mismatch / missing map
# ---------------------------------------------------------------------------

class TestApprovalMapValidation:
    def test_missing_map_returns_none(self, tmp_path: Path) -> None:
        bundle = load_exact_approval_bundle(tmp_path, "B24_RL2")
        assert bundle is None

    def test_form_id_mismatch_returns_none(self, tmp_path: Path) -> None:
        (tmp_path / "schemas" / "maps").mkdir(parents=True)
        (tmp_path / "schemas" / "maps" / "B24_RL2.json").write_text(
            json.dumps({"form_id": "WRONG", "form_version": "1", "template_path": "t.docx", "fields": {"x": {"field_id": "x", "table_index": 0, "row": 0, "col": 0}}}),
            encoding="utf-8",
        )
        assert load_exact_approval_bundle(tmp_path, "B24_RL2") is None

    def test_version_mismatch_returns_none(self, tmp_path: Path) -> None:
        (tmp_path / "schemas" / "maps").mkdir(parents=True)
        (tmp_path / "schemas" / "templates").mkdir(parents=True)
        (tmp_path / "templates").mkdir(parents=True)
        (tmp_path / "templates" / "t.docx").write_bytes(b"PK\x05\x06" + b"\x00" * 18)
        (tmp_path / "schemas" / "templates" / "m.json").write_text(
            json.dumps(
                {
                    "template": "t.docx",
                    "cells": [{"field_id": "x", "table_index": 0, "row": 0, "col": 0, "label": "y"}],
                }
            ),
            encoding="utf-8",
        )
        (tmp_path / "schemas" / "maps" / "B24_RL2.json").write_text(
            json.dumps({
                "form_id": "B24_RL2",
                "form_version": "2026",
                "manifest_path": "schemas/templates/m.json",
                "template_path": "templates/t.docx",
                "fields": {"x": {"field_id": "x", "table_index": 0, "row": 0, "col": 0}},
            }),
            encoding="utf-8",
        )
        result = load_exact_approval_bundle_checked(tmp_path, "B24_RL2", expected_version="2025")
        assert result.bundle is None
        assert any("version mismatch" in e for e in result.errors)


    def test_missing_template_path_rejects_bundle(self, tmp_path: Path) -> None:
        (tmp_path / "schemas" / "maps").mkdir(parents=True)
        (tmp_path / "schemas" / "templates").mkdir(parents=True)
        (tmp_path / "schemas" / "templates" / "m.json").write_text(
            json.dumps(
                {
                    "template": "t.docx",
                    "cells": [{"field_id": "x", "table_index": 0, "row": 0, "col": 0, "label": "y"}],
                }
            ),
            encoding="utf-8",
        )
        (tmp_path / "schemas" / "maps" / "B24_RL2.json").write_text(
            json.dumps({
                "form_id": "B24_RL2",
                "form_version": "2026",
                "manifest_path": "schemas/templates/m.json",
                "template_path": "templates/missing.docx",
                "fields": {"x": {"field_id": "x", "table_index": 0, "row": 0, "col": 0}},
            }),
            encoding="utf-8",
        )
        result = load_exact_approval_bundle_checked(tmp_path, "B24_RL2")
        assert result.bundle is None
        assert any("template not found" in e for e in result.errors)

    def test_duplicate_coordinates_reported(self, tmp_path: Path) -> None:
        (tmp_path / "schemas" / "maps").mkdir(parents=True)
        (tmp_path / "templates").mkdir(parents=True)
        (tmp_path / "templates" / "t.docx").write_bytes(b"PK\x05\x06" + b"\x00" * 18)
        (tmp_path / "schemas" / "maps" / "B24_RL2.json").write_text(
            json.dumps({
                "form_id": "B24_RL2",
                "form_version": "2026",
                "template_path": "templates/t.docx",
                "fields": {
                    "a": {"field_id": "a", "table_index": 0, "row": 1, "col": 0},
                    "b": {"field_id": "b", "table_index": 0, "row": 1, "col": 0},
                },
            }),
            encoding="utf-8",
        )
        result = load_exact_approval_bundle_checked(tmp_path, "B24_RL2")
        assert result.bundle is None
        assert any("duplicate coordinate" in e for e in result.errors)


# ---------------------------------------------------------------------------
# No fallback mapping regression
# ---------------------------------------------------------------------------

class TestNoFallbackMapping:
    def test_unknown_form_has_no_map(self) -> None:
        root = _repo_root()
        assert load_exact_approval_bundle(root, "B91") is None

    def test_no_generic_or_nearest_map_fallback(self, tmp_path: Path) -> None:
        """Map loader must not return a bundle for a form that has no file."""
        (tmp_path / "schemas" / "maps").mkdir(parents=True)
        (tmp_path / "schemas" / "maps" / "B24_RL2.json").write_text(
            json.dumps({"form_id": "B24_RL2", "form_version": "1", "template_path": "t.docx", "fields": {"x": {"field_id": "x", "table_index": 0, "row": 0, "col": 0}}}),
            encoding="utf-8",
        )
        assert load_exact_approval_bundle(tmp_path, "B81") is None

    def test_canonical_map_file_only_legacy_duplicate_ignored(self, tmp_path: Path) -> None:
        root = tmp_path
        (root / "schemas" / "maps").mkdir(parents=True)
        (root / "schemas" / "templates").mkdir(parents=True)
        (root / "templates").mkdir(parents=True)
        manifest = {"template": "t.docx", "cells": [{"field_id": "f", "table_index": 0, "row": 1, "col": 0, "label": "x"}]}
        (root / "schemas" / "templates" / "M.json").write_text(json.dumps(manifest), encoding="utf-8")
        (root / "templates" / "t.docx").write_bytes(b"PK\x05\x06" + b"\x00" * 18)

        canonical = {"form_id": "B24_RL2", "form_version": "1", "manifest_path": "schemas/templates/M.json", "template_path": "templates/t.docx", "fields": {"f": {"field_id": "f", "table_index": 0, "row": 1, "col": 0}}}
        legacy = {"form_id": "B24_RL2", "template": "wrong.docx", "fields": {"g": {"field_id": "g", "table_index": 0, "row": 0, "col": 0}}}
        (root / "schemas" / "maps" / "B24_RL2.json").write_text(json.dumps(canonical), encoding="utf-8")
        (root / "schemas" / "maps" / "B24_RL2.approval_map.json").write_text(json.dumps(legacy), encoding="utf-8")

        bundle = load_exact_approval_bundle(root, "B24_RL2")
        assert bundle is not None
        assert bundle.map_path.name == "B24_RL2.json"
        assert bundle.approval_map["fields"]["f"]["field_id"] == "f"

    def test_manifest_coordinate_mismatch_rejects_bundle(self, tmp_path: Path) -> None:
        (tmp_path / "schemas" / "maps").mkdir(parents=True)
        (tmp_path / "schemas" / "templates").mkdir(parents=True)
        (tmp_path / "templates").mkdir(parents=True)
        (tmp_path / "templates" / "t.docx").write_bytes(b"PK\x05\x06" + b"\x00" * 18)
        (tmp_path / "schemas" / "templates" / "bad.json").write_text(
            json.dumps(
                {
                    "template": "t.docx",
                    "cells": [{"field_id": "x", "table_index": 0, "row": 9, "col": 9, "label": "nope"}],
                }
            ),
            encoding="utf-8",
        )
        (tmp_path / "schemas" / "maps" / "B24_RL2.json").write_text(
            json.dumps({
                "form_id": "B24_RL2",
                "form_version": "2026",
                "manifest_path": "schemas/templates/bad.json",
                "template_path": "templates/t.docx",
                "fields": {"x": {"field_id": "x", "table_index": 0, "row": 0, "col": 0}},
            }),
            encoding="utf-8",
        )
        result = load_exact_approval_bundle_checked(tmp_path, "B24_RL2")
        assert result.bundle is None
        assert any("does not match" in e for e in result.errors)


# ---------------------------------------------------------------------------
# Structure guard handoff gate
# ---------------------------------------------------------------------------

class TestStructureGuardHandoff:
    def test_guard_pass_produces_filled_docx(self, tmp_path: Path) -> None:
        root = _repo_root()
        template = root / "templates" / B24_SHARED_TEMPLATE_DOCX
        if not template.is_file():
            pytest.skip("missing template")

        inbox = tmp_path / "inbox"
        inbox.mkdir()
        (inbox / "evidence.txt").write_text(
            "Cover Page Facility: Midwest Tank Rail Inc\nB24 RL2 objective evidence Date: 2026-05-07\n"
            "B81 stub sill evidence Car: DOTX 123456\nB89 insulation test plate\nB90 RLS return to service Auditor: Casey",
            encoding="utf-8",
        )
        result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run")
        manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
        guard = json.loads((tmp_path / "run" / "structure_guard_report.json").read_text(encoding="utf-8"))
        assert guard["pass"] is True
        assert manifest["structure_guard_passed"] is True
        assert result.filled_docx_path is not None
        assert result.filled_docx_path.is_file()

    def test_cover_page_smoke_values_land_in_separate_cells(self, tmp_path: Path) -> None:
        root = _repo_root()
        inbox = tmp_path / "inbox"
        inbox.mkdir()
        (inbox / "evidence.txt").write_text(
            "Cover Page Facility: Midwest Tank Rail Inc\n"
            "B24 RL2 objective evidence Date: 2026-05-07\n",
            encoding="utf-8",
        )

        run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run")

        doc = Document(tmp_path / "run" / "filled" / "Cover_Page_filled.docx")
        facility_cell = doc.tables[0].rows[4].cells[0].text
        date_cell = doc.tables[0].rows[4].cells[3].text
        assert facility_cell == "Midwest Tank Rail Inc"
        assert date_cell == "2026-05-07"

    def test_guard_fail_discards_filled_docx(self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
        root = _repo_root()
        template = root / "templates" / B24_SHARED_TEMPLATE_DOCX
        if not template.is_file():
            pytest.skip("missing template")

        inbox = tmp_path / "inbox"
        inbox.mkdir()
        (inbox / "evidence.txt").write_text("Cover Page Facility: Midwest\nB24 RL2 evidence Date: 2026-05-07\n", encoding="utf-8")

        from b2_automation.ooxml_writer import PatchOutcome

        def _fail_patch(*args, **kwargs):
            outp = kwargs.get("output_path")
            if outp is None and len(args) > 3:
                outp = args[3]
            return PatchOutcome(
                output_docx=Path(outp).resolve(),
                structure_guard_report=None,
                structure_guard_passed=False,
                patched_fields=(),
                errors=("injected guard failure",),
            )

        monkeypatch.setattr("b2_automation.inbox_pipeline.patch_docx_cells", _fail_patch)

        result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run")
        assert result.status == "review_required"

        manifest = json.loads(result.manifest_path.read_text(encoding="utf-8"))
        assert manifest["structure_guard_passed"] is False
        failed_forms = manifest.get("structure_guard_failed_forms") or []
        assert len(failed_forms) >= 1

        review = json.loads(result.review_json_path.read_text(encoding="utf-8"))
        assert review.get("structure_guard_failed_forms")
        discard_detail = review.get("structure_guard_discard_detail") or []
        assert len(discard_detail) >= 1

        filled_dir = tmp_path / "run" / "filled"
        filled_files = list(filled_dir.glob("*_filled.docx")) if filled_dir.is_dir() else []
        assert filled_files == [], "No filled DOCX should survive a structure guard failure"

    def test_missing_map_produces_review_only(self, tmp_path: Path) -> None:
        fake_root = tmp_path / "fake_root"
        (fake_root / "schemas" / "maps").mkdir(parents=True)
        (fake_root / "templates").mkdir(parents=True)

        inbox = tmp_path / "inbox"
        inbox.mkdir()
        (inbox / "evidence.txt").write_text("B24 RL2 evidence Facility: Midwest Date: 2026-01-01", encoding="utf-8")
        result = run_inbox_pipeline(root=fake_root, inbox=inbox, out_dir=tmp_path / "run")
        assert result.review_json_path.is_file()
        review = json.loads(result.review_json_path.read_text(encoding="utf-8"))

        filled_dir = tmp_path / "run" / "filled"
        filled_files = list(filled_dir.glob("*_filled.docx")) if filled_dir.is_dir() else []
        assert filled_files == [], "No filled DOCX when approval map is missing"

        docx_gen = review.get("docx_generation") or []
        for item in docx_gen:
            assert item.get("filled_docx") is None
