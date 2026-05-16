"""Inbox pipeline regression tests."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pytest
from docx import Document

from b2_automation.docupipe_client import DocuPipeConfigError, process_pdf
from b2_automation.inbox_pipeline import (
    _best_value_for_label,
    _clear_scoped_filled_docx,
    _label_value_is_compatible,
    run_inbox_pipeline,
)
from b2_automation.local_extraction import DEFAULT_REVIEW_FORMS, supported_evidence_files


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _manifest(result: object) -> dict[str, object]:
    return json.loads(result.manifest_path.read_text(encoding="utf-8"))


def _review(result: object) -> dict[str, object]:
    return json.loads(result.review_json_path.read_text(encoding="utf-8"))


def _fixture_text(name: str) -> str:
    fixture = _repo_root() / "tests" / "fixtures" / name
    return fixture.read_text(encoding="utf-8")


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def test_cover_page_does_not_auto_fill_b24_body_rows(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "mixed_cover_b24.txt").write_text(
        "\n".join(
            [
                "Cover Page",
                "Station Stencil/QA Code: DLGA",
                "Audit Type: SP",
                "Open Meeting Date: 2026-05-05",
                "Closing Meeting Date: 2026-05-07",
                "BOE Lead Auditor: Lucho Rodriguez",
                "Tank Car Owner (TCO) Name: CIT",
                "Car Mark and Number: DBUX 250086",
                "AAR Form 4-2 (AAR No.): L016048A",
                "Drawing Number: D43520",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("Cover_Page",))

    manifest = _manifest(result)
    cover_docx = manifest["docx_generation"][0]
    assert cover_docx["form_id"] == "Cover_Page"
    assert cover_docx["auto_table_fields"] == []
    assert cover_docx["auto_table_manual_fields"] == []
    text = _docx_text(Path(cover_docx["filled_docx"]))
    assert "PART 1: General INFORMATiON" in text
    assert "Station Stencil/QA Code: DLGA" in text
    assert "Open Meeting Date: 2026-05-05" in text
    assert "CAR OWNER PERMISSIONS" not in text
    assert "DBUX 250086" not in text
    assert "L016048A" not in text


@pytest.mark.parametrize(
    ("label", "value"),
    [
        ("Date Approved", "day where the action"),
        ("Equipment ID", "equipment Red equipment Red equipment Red"),
        ("Welder ID", "121405)"),
        ("Location", "1J1~'"),
        ("Material Description", "Confirmacion por correo electronico"),
        ("Facility Location", "T-joint"),
        ("Technician", "a) a) a)"),
    ],
)
def test_auto_table_rejects_known_ocr_junk(label: str, value: str) -> None:
    assert _label_value_is_compatible(label, value) is False


def test_auto_table_allows_contextual_confirmation_and_stub_type_values() -> None:
    assert _label_value_is_compatible("Written Instructions from TCO", "Confirmacion por correo electronico") is True
    assert _label_value_is_compatible("Stub Sill Type", "T-joint") is True


def test_auto_table_does_not_use_generic_control_text_for_thermocouples() -> None:
    evidence = {
        "control": [
            "Action type + Facility + Date + Consecutive. The Date: day where the action is uploaded.",
        ]
    }

    assert _best_value_for_label("Control thermocouples required procedure", evidence) is None


def _assert_auto_table_manual_tracked_separate(manifest: dict[str, object]) -> None:
    manual = manifest.get("manual_fields") or {}
    auto_manual = manifest.get("auto_table_manual_fields") or {}
    docx_gen = manifest.get("docx_generation") or []
    for docx_run in docx_gen:
        if not isinstance(docx_run, dict) or docx_run.get("status") != "filled":
            continue
        form_id = str(docx_run.get("form_id", ""))
        run_manual = set(docx_run.get("manual_fields") or [])
        run_auto_manual = set(docx_run.get("auto_table_manual_fields") or [])

        intersection = run_manual & run_auto_manual
        assert not intersection, f"Fields in both manual and auto_table_manual: {intersection}"

        for f in run_auto_manual:
            assert str(f).startswith("auto_table."), f"auto_table_manual_fields {f} must start with auto_table."
        for f in run_manual:
            assert not str(f).startswith("auto_table."), f"manual_fields {f} must not start with auto_table."

        assert set(manual.get(form_id) or []) == run_manual
        assert set(auto_manual.get(form_id) or []) == run_auto_manual


def test_inbox_pipeline_local_default_generates_all_form_packets(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "packet_one.txt").write_text(
        "\n".join(
            [
                "Cover Page Facility: Midwest Tank Rail Inc",
                "B24 RL2 objective evidence Date: 2026-05-07",
                "B81 stub sill evidence Car: DOTX 123456",
                "B89 insulation test plate evidence",
                "B90 RLS return to service evidence Auditor: Casey",
            ]
        ),
        encoding="utf-8",
    )
    (inbox / "evidence_sample.txt").write_text(
        "Sample audit evidence line for local inbox dry run. Company: Example Railroad Year: 2025",
        encoding="utf-8",
    )
    (inbox / "procedure.txt").write_text(
        "Cover Page procedure note. The Facility: assigned code for AAR and internal action tracking.",
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run")

    sample_tpl = root / "templates" / "B24_RL2.docx"
    run_manifest_data = _manifest(result)
    assert result.status in {"success", "review_required"}
    if sample_tpl.is_file():
        guard = json.loads((tmp_path / "run" / "structure_guard_report.json").read_text(encoding="utf-8"))
        assert guard["pass"] is True
        assert run_manifest_data.get("structure_guard_passed") is True
        assert result.filled_docx_path is not None and result.filled_docx_path.is_file()
        assert result.filled_docx_path.name.endswith("_filled.docx")
        assert result.filled_docx_paths
        assert all(p.is_file() for p in result.filled_docx_paths)
        assert run_manifest_data.get("structure_guard_failed_forms") == []
    else:
        assert result.filled_docx_path is None
        assert result.filled_docx_paths == ()
    assert result.manifest_path.is_file()
    assert result.review_json_path.is_file()
    assert result.review_md_path.is_file()

    run_dir = tmp_path / "run"
    assert not list((run_dir / "raw").glob("*.docupipe.json"))
    assert (run_dir / "raw" / "packet_one.ocr.json").is_file()
    assert (run_dir / "raw" / "packet_one.metadata.json").is_file()
    assert (run_dir / "raw" / "packet_one.chunks.json").is_file()
    assert (run_dir / "raw" / "local_rag_retrieval.json").is_file()
    assert (run_dir / "rag_selection_report.json").is_file()

    review = _review(result)
    manifest = _manifest(result)
    assert review["docupipe_used"] is False
    assert manifest["docupipe_used"] is False
    assert manifest["legacy_adapter_used"] is False
    assert set(review["forms"]) == set(DEFAULT_REVIEW_FORMS)
    assert review["production_scope_forms"] == list(DEFAULT_REVIEW_FORMS)
    assert "evidence_sample.txt" not in {row["source_file"] for row in review["inputs"]}
    canonical = json.loads((run_dir / "canonical_evidence.json").read_text(encoding="utf-8"))
    facility_values = [
        str(row.get("selected_value") or "")
        for row in canonical["fields"]
        if row.get("field_id") == "facility_name" and row.get("selected_value")
        and not str(row.get("selected_value") or "").startswith("REVIEW_REQUIRED")
    ]
    assert facility_values
    assert set(facility_values) == {"Midwest Tank Rail Inc"}
    cover_packet = json.loads((run_dir / "review" / "Cover_Page_evidence_packet.json").read_text(encoding="utf-8"))
    assert not [
        row
        for row in cover_packet["field_suggestions"]
        if row["field_id"] == "facility_name" and row["candidate_value"] == "assigned code for"
    ]
    for form in DEFAULT_REVIEW_FORMS:
        packet_path = run_dir / "review" / f"{form}_evidence_packet.json"
        assert packet_path.is_file()
        assert (run_dir / "review" / f"{form}_review.md").is_file()
        packet = json.loads(packet_path.read_text(encoding="utf-8"))
        assert packet.get("production_scope") is True
        assert "b24_rl1_legacy_only" not in packet
        assert "missing_fields" in packet
        assert "conflicts" in packet
        assert "low_confidence_fields" in packet
        assert "decision_summary" in packet
        assert "field_decisions" in packet
        for row in packet["field_decisions"]:
            assert row["state"] in {
                "FILL",
                "BLANK",
                "REVIEW_REQUIRED",
                "MISSING",
                "CONFLICT",
                "LOW_CONFIDENCE",
            }
        assert packet["field_suggestions"]
    _assert_auto_table_manual_tracked_separate(manifest)


def test_clear_scoped_filled_docx_removes_stale_outputs(tmp_path: Path) -> None:
    filled = tmp_path / "filled"
    filled.mkdir()
    stale = filled / "B81_filled.docx"
    stale.write_bytes(b"PK\x03\x04")
    untouched = filled / "B89_filled.docx"
    untouched.write_bytes(b"PK\x03\x04")
    _clear_scoped_filled_docx(filled, ("B81",))
    assert not stale.is_file()
    assert untouched.is_file()


def test_inbox_pipeline_fills_b81_with_manual_markers_when_required_values_are_missing(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b81_noise.txt").write_text(
        "\n".join(
            [
                "B81 stub sill evidence and side sill review for tank car repair.",
                "Facility: Demo Rail Shop",
                "This packet intentionally omits the required date.",
            ]
        ),
        encoding="utf-8",
    )
    out_dir = tmp_path / "run"
    filled_dir = out_dir / "filled"
    filled_dir.mkdir(parents=True)
    stale = filled_dir / "B81_filled.docx"
    stale.write_bytes(b"PK\x03\x04")

    result = run_inbox_pipeline(
        root=root,
        inbox=inbox,
        out_dir=out_dir,
        review_forms=("B81",),
        low_confidence_threshold=0.0,
    )

    manifest = _manifest(result)
    review = _review(result)
    b81_docx = next(item for item in manifest["docx_generation"] if item["form_id"] == "B81")
    b81_decisions = review["form_packets"]["B81"]["field_decisions"]

    assert result.status == "review_required"
    assert any(row["state"] == "FILL" for row in b81_decisions)
    assert b81_docx["status"] == "filled"
    assert b81_docx["filled_docx"] is not None
    assert Path(str(b81_docx["filled_docx"])).is_file()
    assert manifest["review_blocked_forms"] == []
    assert manifest["skipped_review_required"] == []
    assert "B81" in manifest["manual_fields"]
    assert "car.mark" in manifest["manual_fields"]["B81"]
    assert stale.exists()
    assert stale.stat().st_size > 4
    assert "REVIEW_REQUIRED" in _docx_text(stale)
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_aliases_car_number_to_b81_car_mark(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b81_car_number.txt").write_text(
        "\n".join(
            [
                "B81 stub sill evidence for tank car repair.",
                "Facility: Demo Rail Shop",
                "Date: 2026-05-07",
                "Car: DOTX 123456",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B81",))

    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B81"]["field_decisions"]}
    assert selected["car.mark"] == "DOTX 123456"
    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert "car.mark" in docx["patched_fields"]
    filled_doc = Document(str(result.filled_docx_path))
    assert "DOTX 123456" in filled_doc.tables[0].rows[6].cells[0].text


def test_inbox_pipeline_fills_b81_docx_when_only_basic_fields_are_present(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b81_bilingual.txt").write_text(
        "\n".join(
            [
                "B81 stub sill evidence for tank car repair.",
                "Estacion/ Station:",
                "Taller Mexico FTVM",
                "Fecha I Date: 06-Mayo-2025",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B81",))

    review = _review(result)
    packet = review["form_packets"]["B81"]
    selected = {row["field_id"]: row["selected_value"] for row in packet["field_decisions"]}
    assert selected["facility_name"] == "Taller Mexico FTVM"
    assert selected["date"] == "2025-05-06"
    assert packet["missing_fields"] == []
    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert result.status == "review_required"
    assert docx["status"] == "filled"
    assert result.filled_docx_path is not None
    assert result.filled_docx_paths
    assert "car.mark" in manifest["manual_fields"]["B81"]
    assert "REVIEW_REQUIRED" in _docx_text(result.filled_docx_path)
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_fills_b81_run_level_evidence_with_manual_markers(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "Adobe Scan May 05, 2026.txt").write_text(
        "B81 stub sill evidence for tank car repair.",
        encoding="utf-8",
    )
    (inbox / "b24.1.txt").write_text(
        "B24 RL2 evidence\nEstacion/ Station: Taller Mexico FTVM\n",
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B81",))

    review = _review(result)
    packet = review["form_packets"]["B81"]
    selected = {row["field_id"]: row["selected_value"] for row in packet["field_decisions"]}
    assert selected["facility_name"] == "Taller Mexico FTVM"
    assert selected["date"] == "2026-05-05"
    assert packet["missing_fields"] == []
    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert result.status == "review_required"
    assert docx["status"] == "filled"
    assert result.filled_docx_path is not None
    assert "B81" in manifest["manual_fields"]
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_patches_multiple_b24_mapped_fields(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_mapped.txt").write_text(
        "\n".join(
            [
                "B24 RL2 MANTENIMIENTO Y MODIFICACION DE LOS CARROS TANQUE RL2",
                "Estacion/ Station: Taller Mexico FTVM",
                "Fecha I Date: 06-Mayo-2025",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission/Instruction Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Mark: PROBETA MUESTRA",
                "Car Number: PAWCT-824",
                "Car Type: TANK CAR",
                "pitp_document_name: PC-TC-01",
                "pitp_id: PC-TC-01",
                "pitp_approved_by: Casey",
                "pitp_date_approved: 2026-05-05",
                "aar_form_4_2_number: AAR-42-001",
                "four_two_drawing_number: DWG-100",
                "Specimen plate A516 Grado 70",
                "test_plate_tank_mtr: MTR-777",
                "attachment_material: A36",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert result.status in {"success", "review_required"}
    assert {
        "facility_name",
        "tco_permission_date",
        "tco_written_instructions",
        "car_mark",
        "tank_design_spec",
        "test_plate_tank_material",
    }.issubset(set(docx["patched_fields"]))
    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B24_RL2"]["field_decisions"]}
    assert selected["car_mark"] == "PAWCT-824"
    assert "tank_design_spec" not in selected or str(selected["tank_design_spec"]).startswith("REVIEW_REQUIRED")
    assert "tank_design_spec" in manifest["manual_fields"]["B24_RL2"]
    assert selected["facility_name"] == "CIT"
    assert selected["tco_written_instructions"] == "Repair authorized per customer email"
    filled_doc = Document(str(result.filled_docx_path))
    assert "Repair authorized" in filled_doc.tables[0].rows[4].cells[8].text
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_does_not_use_b24_activity_title_as_tco_instructions(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_scope_only.txt").write_text(
        "\n".join(
            [
                "B24 RL2 MANTENIMIENTO Y MODIFICACION DE LOS CARROS TANQUE RL2",
                "Estacion/ Station: Taller Mexico FTVM",
                "Fecha I Date: 06-Mayo-2025",
                "Car Mark: PROBETA MUESTRA",
                "Car Number: PAWCT-824",
                "Car Type: TANK CAR",
                "pitp_document_name: PC-TC-01",
                "pitp_id: PC-TC-01",
                "pitp_approved_by: Casey",
                "pitp_date_approved: 2026-05-05",
                "aar_form_4_2_number: AAR-42-001",
                "four_two_drawing_number: DWG-100",
                "Specimen plate A516 Grado 70",
                "test_plate_tank_mtr: MTR-777",
                "attachment_material: A36",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert "tco_written_instructions" in docx["manual_fields"]
    filled_doc = Document(str(result.filled_docx_path))
    instruction_cell = filled_doc.tables[0].rows[4].cells[8].text
    assert "MANTENIMIENTO Y MODIFICACION" not in instruction_cell
    assert "REVIEW_REQUIRED: tco_written_instructions" in instruction_cell
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_requires_tco_label_for_email_confirmation_instructions(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_email_confirmation.txt").write_text(
        "\n".join(
            [
                "B24 RL2 evidence",
                "Confirmaci6n por correo electr6nico",
                "Car Mark: PROBETA MUESTRA",
                "Car Number: PAWCT-824",
                "Car Type: TANK CAR",
                "PC-TC-01 Alondra Navarro c-4--nov---21 Primera edicion",
                "aar_form_4_2_number: AAR-42-001",
                "four_two_drawing_number: DWG-100",
                "Specimen plate A516 Grado 70",
                "test_plate_tank_mtr: MTR-777",
                "attachment_material: A36",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    manifest = _manifest(result)
    assert "tco_written_instructions" in manifest["manual_fields"]["B24_RL2"]
    filled_doc = Document(str(result.filled_docx_path))
    assert "Confirmacion por correo electronico" not in filled_doc.tables[0].rows[4].cells[8].text
    assert "REVIEW_REQUIRED: tco_written_instructions" in filled_doc.tables[0].rows[4].cells[8].text


def test_inbox_pipeline_recovers_b24_ocr_pitp_date_and_material(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_ocr.txt").write_text(
        "\n".join(
            [
                "B24 RL2 evidence",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission/Instruction Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Mark: PROBETA MUESTRA",
                "Car Number: PAWCT-824",
                "Car Type: TANK CAR",
                "PC-TC-01 Alondra Navarro c-4--nov---21 Primera edicion",
                "aar_form_4_2_number: AAR-42-001",
                "four_two_drawing_number: DWG-100",
                "Specimen plate A51G Grado 70",
                "test_plate_tank_mtr: MTR-777",
                "attachment_material: A36",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B24_RL2"]["field_decisions"]}
    assert selected["pitp_date_approved"] == "2021-11-04"
    assert selected["test_plate_tank_material"] == "A516 Gr. 70"


def test_inbox_pipeline_recovers_manual_b24_style_fields_and_rejects_junk_dates(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "manual_b24_style.txt").write_text(
        "\n".join(
            [
                "B24 RL2 evidence",
                "TCO Name: CIT",
                "TCO Permission Date: 5/20/2024",
                "Written Instructions from TCO: Facility received written confirmation from owner Allowing FACILITY Procedures.",
                "Date Approved: day where the action",
                "PITP: PITP / PC-TC-01 / A Navarre / 11/4/2021 / 0",
                "Car Mark: DBUX 250086",
                "Design Spec: DOT111A100W1",
                "Stencil Spec: AAR211A100W1",
                "AAR Form 4-2: L016048A",
                "Drawing: D43520",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B24_RL2"]["field_decisions"]}
    assert selected["facility_name"] == "CIT"
    assert selected["tco_permission_date"] == "5/20/2024"
    assert selected["tco_written_instructions"] == "Facility received written confirmation from owner Allowing FACILITY Procedures"
    assert selected["pitp_document_name"] == "PITP"
    assert selected["pitp_id"] == "PC-TC-01"
    assert selected["pitp_approved_by"] == "A Navarre"
    assert selected["pitp_date_approved"] == "11/4/2021"
    assert selected["pitp_rev"] == "0"
    assert selected["car_mark"] == "DBUX 250086"
    assert selected["tank_design_spec"] == "DOT111A100W1 / AAR211A100W1"
    assert selected["aar_form_4_2_number"] == "L016048A"
    assert selected["four_two_drawing_number"] == "D43520"
    assert all("day where the action" not in str(value).lower() for value in selected.values())

    filled_doc = Document(str(result.filled_docx_path))
    assert "day where the action" not in _docx_text(result.filled_docx_path).lower()
    assert "DOT111A100W1 / AAR211A100W1" in filled_doc.tables[0].rows[9].cells[2].text


def test_inbox_pipeline_validates_b24_table_pair_fixture_end_to_end(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "DLGA-B24_table_pairs.txt").write_text(_fixture_text("dlga_b24_table_pairs_fixture.txt"), encoding="utf-8")

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    manifest = _manifest(result)
    b24_docx = next(item for item in manifest["docx_generation"] if item["form_id"] == "B24_RL2")
    assert b24_docx["status"] == "filled"
    assert b24_docx["structure_guard_passed"] is True
    assert manifest["structure_guard_passed"] is True
    assert (tmp_path / "run" / "structure_guard_report.json").is_file()

    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B24_RL2"]["field_decisions"]}
    assert selected["facility_name"] == "CIT"
    assert selected["tco_permission_date"] == "5/20/2024"
    assert selected["tco_written_instructions"] == "Facility received written confirmation from owner Allowing FACILITY Procedures"
    assert selected["pitp_document_name"] == "PITP"
    assert selected["pitp_id"] == "PC-TC-01"
    assert selected["pitp_approved_by"] == "A Navarre"
    assert selected["pitp_date_approved"] == "11/4/2021"
    assert selected["pitp_rev"] == "0"
    assert selected["tank_design_spec"] == "DOT111A100W1 / AAR211A100W1"
    assert selected["aar_form_4_2_number"] == "L016048A"
    assert selected["four_two_drawing_number"] == "D43520"
    assert selected["tco.name"] == "CIT"

    manual_map_fields = set(manifest["manual_fields"]["B24_RL2"])
    assert manual_map_fields == {"test_plate_tank_material", "test_plate_tank_mtr", "attachment_material"}
    assert result.status == "review_required"

    text = _docx_text(result.filled_docx_path).lower()
    for junk in ("day where the action", "a):", "malformed ocr fragment"):
        assert junk not in text
    for should_not_be_manual in (
        "review_required: facility_name",
        "review_required: tco_permission_date",
        "review_required: pitp_document_name",
        "review_required: pitp_id",
        "review_required: tank_design_spec",
        "review_required: four_two_drawing_number",
    ):
        assert should_not_be_manual not in text


def test_inbox_pipeline_b24_aliases_prefer_tco_name_over_facility_alias_noise(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_alias_noise.txt").write_text(
        "\n".join(
            [
                "B24 RL2 evidence",
                "facility_name: WRONG FACILITY VALUE",
                "TCO Name: CIT",
                "TCO Permission Date: 5/20/2024",
                "Written Instructions from TCO: Facility received written confirmation from owner Allowing FACILITY Procedures.",
                "PITP: PITP / PC-TC-01 / A Navarre / 11/4/2021 / 0",
                "Car Mark: DBUX 250086",
                "Design Spec: DOT111A100W1",
                "Stencil Spec: AAR211A100W1",
                "AAR Form 4-2: L016048A",
                "Drawing: D43520",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))
    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B24_RL2"]["field_decisions"]}
    assert selected["facility_name"] == "CIT"
    assert selected["tco.name"] == "CIT"


def test_inbox_pipeline_b24_aliases_do_not_hardcode_customer_name(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_alt_tco.txt").write_text(
        "\n".join(
            [
                "B24 RL2 evidence",
                "TCO Name: GATX",
                "TCO Permission Date: 5/20/2024",
                "Written Instructions from TCO: Owner approved procedures in writing.",
                "PITP: PITP / PC-TC-77 / J Smith / 11/4/2021 / 0",
                "Car Mark: DBUX 250086",
                "Design Spec: DOT111A100W1",
                "Stencil Spec: AAR211A100W1",
                "AAR Form 4-2: L016048A",
                "Drawing: D43520",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))
    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B24_RL2"]["field_decisions"]}
    assert selected["facility_name"] == "GATX"
    assert selected["tco.name"] == "GATX"
    assert selected["pitp_document_name"] == "PITP"


def test_inbox_pipeline_b24_does_not_fill_numeric_or_date_targets_from_prose_only_labels(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_prose_noise.txt").write_text(
        "\n".join(
            [
                "[structured_docx_table_evidence]",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission/Instruction Received from TCO: 5/20/2024",
                "Written Instructions from TCO: Facility received written confirmation from owner Allowing FACILITY Procedures.",
                "PITP Document Name: PITP",
                "PITP ID: should follow owner guidance in procedure",
                "Date Approved: day where the action",
                "AAR Form 4-2 (AAR No.): see owner email thread",
                "Drawing Number: as noted by inspector in general prose",
                "Car Mark and Number: DBUX 250086",
                "Tank Car Design Spec/Stencil Spec: DOT111A100W1 / AAR211A100W1",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))
    manifest = _manifest(result)
    manual_fields = set(manifest["manual_fields"]["B24_RL2"])
    assert {"pitp_id", "pitp_date_approved", "aar_form_4_2_number", "four_two_drawing_number"} <= manual_fields
    text = _docx_text(result.filled_docx_path).lower()
    for junk in (
        "day where the action",
        "should follow owner guidance in procedure",
        "see owner email thread",
        "as noted by inspector in general prose",
    ):
        assert junk not in text


def test_inbox_pipeline_does_not_autofill_generic_date_approved_from_pitp_date(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b24_ocr.txt").write_text(
        "\n".join(
            [
                "B24 RL2 evidence",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission/Instruction Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Mark: PROBETA MUESTRA",
                "Car Number: PAWCT-824",
                "Car Type: TANK CAR",
                "PC-TC-01 Alondra Navarro c-4--nov---21 Primera edicion",
                "aar_form_4_2_number: AAR-42-001",
                "four_two_drawing_number: DWG-100",
                "Specimen plate A516 Grado 70",
                "test_plate_tank_mtr: MTR-777",
                "attachment_material: A36",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    filled_doc = Document(str(result.filled_docx_path))
    assert "2021-11-04" not in filled_doc.tables[0].rows[11].cells[16].text
    manifest = _manifest(result)
    b24_docx = manifest["docx_generation"][0]
    assert "auto_table.t0.r11.c16.date_approved" in b24_docx["auto_table_manual_fields"]


def test_inbox_pipeline_patches_b89_from_docupipe_schema_json(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "Adobe Scan May 05, 2026.json").write_text(
        json.dumps(
            {
                "activity": {
                    "code": "B89",
                    "repairLevel": "RLJ",
                    "scope": "MANTENIMIENTO, MODIFICACION y CALIFICACION DE SISTEMAS DE SEGURIDAD",
                },
                "demonstration": {
                    "carNumber": "PAWCT-RLJ",
                    "carType": "TANK CAR",
                    "station": "Taller Mexico FTVM",
                },
                "tco": {
                    "name": "CIT",
                    "permissionDate": "2026-05-04",
                    "instructions": "Repair authorized per customer email",
                },
                "aar": {"form42Number": "AAR-89-001"},
                "jacketPatch": {
                    "specimenPlate": "A36 1/8 in o A1110 cal. 11",
                    "patchPlateSize": "12 in X 12 in",
                    "targetFilletSize": "3/16 in filete",
                },
                "pitp": "PC-TC-01",
            }
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B89",))

    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert result.status in {"success", "review_required"}
    assert docx["status"] == "filled"
    assert {
        "tco.name",
        "tco.permission_date",
        "tco.instructions",
        "car.mark",
        "car.design_spec",
        "safety_system.type",
        "aar.form_4_2.number",
        "materials.insulation.spec",
        "materials.jacket.spec",
        "test_fixture.patch_plate.size",
        "test_fixture.weld.length",
        "pitp.name",
        "pitp.id",
    }.issubset(set(docx["patched_fields"]))
    filled_doc = Document(str(result.filled_docx_path))
    assert "Repair authorized" in filled_doc.tables[0].rows[2].cells[5].text
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_patches_b90_from_stub_sill_packet_text(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b90_packet.txt").write_text(
        "\n".join(
            [
                "B90 Maintenance, Alteration, and Qualification of Tank Car Stub Sills",
                "Estacion/ Station: Taller Mexico FTVM",
                "Fecha I Date: 06-Mayo-2025",
                "MANTENIMIENTO Y MODIFICACION DE STUB SILLS",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Mark: PROBETA MUESTRA",
                "Car Number: PAWCT-B90",
                "Tank Car Design Specification: DOT111A100W1",
                "Stencil Specification: AAR211A100W1",
                "Car Type: TANK CAR",
                "General Arrangement - D-41759 L056040A",
                "Specimen plate A572 Grado 50",
                "Junta de soldaura en T/T-weld joint",
                "Plan de control PC-TC-01",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B90",))

    manifest = _manifest(result)
    docx = manifest["docx_generation"][0]
    assert result.status in {"success", "review_required"}
    assert docx["status"] == "filled"
    assert {
        "tco.name",
        "tco.permission_date",
        "tco.instructions",
        "car.mark",
        "car.design_spec",
        "stub_sill.type",
        "aar.form_4_2.number",
        "materials.stub_sill.spec",
        "pitp.name",
        "pitp.id",
        "stub_sill.procedure.id",
    }.issubset(set(docx["patched_fields"]))
    filled_doc = Document(str(result.filled_docx_path))
    assert "Repair authorized" in filled_doc.tables[0].rows[4].cells[5].text
    _assert_auto_table_manual_tracked_separate(manifest)


def test_inbox_pipeline_b90_rejects_b24_specimen_code_from_shared_inventory(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b90_packet.txt").write_text(
        "\n".join(
            [
                "B90 Maintenance, Alteration, and Qualification of Tank Car Stub Sills",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Number: PAWCT-B90",
                "Tank Car Design Specification: DOT111A100W1",
                "Stencil Specification: AAR211A100W1",
                "Specimen plate A572 Grado 50",
                "Plan de control PC-TC-01",
            ]
        ),
        encoding="utf-8",
    )
    (inbox / "GQAP-2.16_preservation.txt").write_text(
        "Identificacion como sigue: codigos asignados: PAWCT-B24, PAWCT-B90, PAWCT-RLJ",
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B90",))

    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B90"]["field_decisions"]}
    assert selected["car.mark"] == "PAWCT-B90"
    text = _docx_text(result.filled_docx_path)
    assert "PAWCT-B90" in text
    assert "PAWCT-B24" not in text


def test_inbox_pipeline_b90_rejects_t_joint_and_tank_car_as_design_specs(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b90_t_joint_noise.txt").write_text(
        "\n".join(
            [
                "B90 Maintenance, Alteration, and Qualification of Tank Car Stub Sills",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Number: PAWCT-B90",
                "Tank Car Design Specification: T-joint",
                "Stencil Specification: T-joint",
                "Car Type: TANK CAR",
                "General Arrangement - D-41759 L056040A",
                "Specimen plate A572 Grado 50",
                "Plan de control PC-TC-01",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B90",))

    manifest = _manifest(result)
    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B90"]["field_decisions"]}
    assert selected["car.mark"] == "PAWCT-B90"
    assert selected.get("car.design_spec") is None or str(selected["car.design_spec"]).startswith("REVIEW_REQUIRED")
    assert "car.design_spec" in manifest["manual_fields"]["B90"]
    text = _docx_text(result.filled_docx_path).lower()
    assert "t-joint" not in text


def test_inbox_pipeline_normalizes_b89_real_ocr_car_and_material_shape(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "b89_real_ocr_shape.txt").write_text(
        "\n".join(
            [
                "B89 insulation test plate evidence",
                "Tank Car Owner (TCO) Name: CIT",
                "Date Permission Received from TCO: 2026-05-04",
                "Written Instructions from TCO: Repair authorized per customer email.",
                "Car Number: PAWCT-RL.J",
                "Tank Car Design Specification: DOT111A100W1",
                "AAR Form 4-2: L056040A",
                "Specimen/Speclmen plate A361/8' oA11 10cal. 11",
                "Patch plate size 12 in X 12 in",
                "3/16 in filete",
                "Plan de control PC-TC-01",
            ]
        ),
        encoding="utf-8",
    )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B89",))

    review = _review(result)
    selected = {row["field_id"]: row["selected_value"] for row in review["form_packets"]["B89"]["field_decisions"]}
    assert selected["car.mark"] == "PAWCT-RLJ"
    assert selected["car.design_spec"] == "DOT111A100W1"
    assert selected["materials.insulation.spec"] == "A36 1/8 in / A1110 cal. 11"
    text = _docx_text(result.filled_docx_path)
    assert "PAWCT-RLJ" in text
    assert "A361/8" not in text


def test_inbox_pipeline_local_rejects_unknown_review_form(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    (inbox / "evidence.txt").write_text("local evidence", encoding="utf-8")
    with pytest.raises(ValueError, match="Unknown review form"):
        run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL1",))


def test_inbox_pipeline_missing_pdf_folder_content_fails_clearly(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "empty_inbox"
    inbox.mkdir()
    with pytest.raises(FileNotFoundError, match="No supported local evidence files found"):
        run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run")


def test_pdf_inbox_ignores_tracked_smoke_evidence_txt(tmp_path: Path) -> None:
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    pdf = inbox / "real.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    (inbox / "evidence.txt").write_text(
        "Cover Page Facility: Smoke Fixture\nB24 RL2 Date: 2099-01-01\n",
        encoding="utf-8",
    )
    (inbox / "evidence_sample.txt").write_text("sample only", encoding="utf-8")

    assert [p.name for p in supported_evidence_files(inbox)] == [pdf.name]


def test_pdf_inbox_ignores_staged_zip_smoke_evidence_txt(tmp_path: Path) -> None:
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    pdf = inbox / "activitys__b24_1.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    (inbox / "evidence_sample__evidence.txt").write_text(
        "Cover Page Facility: Smoke Fixture\nB24 RL2 Date: 2099-01-01\n",
        encoding="utf-8",
    )

    assert [p.name for p in supported_evidence_files(inbox)] == [pdf.name]


def test_inbox_pipeline_extracts_supported_files_from_zip(tmp_path: Path) -> None:
    root = _repo_root()
    inbox = tmp_path / "inbox"
    inbox.mkdir()
    with zipfile.ZipFile(inbox / "evidence_bundle.zip", "w") as zf:
        zf.writestr(
            "nested/b24_mapped.txt",
            "\n".join(
                [
                    "B24 RL2 MANTENIMIENTO Y MODIFICACION DE LOS CARROS TANQUE RL2",
                    "Estacion/ Station: Taller Mexico FTVM",
                    "Fecha I Date: 06-Mayo-2025",
                    "Tank Car Owner (TCO) Name: CIT",
                    "Date Permission/Instruction Received from TCO: 2026-05-04",
                    "Written Instructions from TCO: Repair authorized per customer email.",
                    "Car Mark: PROBETA MUESTRA",
                    "Car Number: PAWCT-824",
                    "Car Type: TANK CAR",
                    "pitp_document_name: PC-TC-01",
                    "pitp_id: PC-TC-01",
                    "pitp_approved_by: Casey",
                    "pitp_date_approved: 2026-05-05",
                    "aar_form_4_2_number: AAR-42-001",
                    "four_two_drawing_number: DWG-100",
                    "Specimen plate A516 Grado 70",
                    "test_plate_tank_mtr: MTR-777",
                    "attachment_material: A36",
                ]
            ),
        )

    result = run_inbox_pipeline(root=root, inbox=inbox, out_dir=tmp_path / "run", review_forms=("B24_RL2",))

    manifest = _manifest(result)
    review = _review(result)
    source_files = {row["source_file"] for row in review["inputs"]}
    assert any("b24_mapped" in str(source) for source in source_files)
    docx = manifest["docx_generation"][0]
    assert docx["status"] == "filled"
    assert {"facility_name", "car_mark", "tank_design_spec"}.issubset(set(docx["patched_fields"]))
    assert result.filled_docx_path is not None and result.filled_docx_path.is_file()


def test_docupipe_live_mode_missing_credentials_fails(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setenv("B2_DOCUPIPE_STUB", "0")
    monkeypatch.delenv("DOCUPIPE_API_KEY", raising=False)
    monkeypatch.delenv("DOCUPIPE_API_URL", raising=False)
    pdf = tmp_path / "input.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    with pytest.raises(DocuPipeConfigError, match="DOCUPIPE_API_KEY"):
        process_pdf(pdf)
