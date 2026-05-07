"""Dump Word table structure for each DOCX in templates/."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from b2_automation.paths import resolve_project_root


def run_discovery(root: Path | None = None, max_rows: int = 30, cell_preview: int = 160) -> list[Path]:
    root = root or resolve_project_root()
    template_dir = root / "templates"
    output_dir = root / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)

    paths = sorted(template_dir.glob("*.docx"))
    written: list[Path] = []
    for template in paths:
        doc = Document(template)
        out = output_dir / f"{template.stem}_table_map.txt"
        with out.open("w", encoding="utf-8") as f:
            f.write(f"Template: {template.name}\n")
            f.write(f"Tables: {len(doc.tables)}\n\n")
            for i, table in enumerate(doc.tables):
                f.write("=" * 100 + "\n")
                f.write(f"TABLE {i}\n")
                f.write(f"Rows: {len(table.rows)}\n")
                f.write(f"Cols: {len(table.rows[0].cells) if table.rows else 0}\n")
                f.write("-" * 100 + "\n")
                for r, row in enumerate(table.rows[:max_rows]):
                    cells = [" ".join(c.text.split())[:cell_preview] for c in row.cells]
                    f.write(f"Row {r}: " + " | ".join(cells) + "\n")
        written.append(out)
    return written
