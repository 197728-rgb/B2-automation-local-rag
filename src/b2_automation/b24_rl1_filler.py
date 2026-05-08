"""Legacy python-docx B24_RL1 filler.

Production-path DOCX writes should use ``b2_automation.ooxml_writer`` so
``word/document.xml`` is patched without full Word reserialization.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.table import _Cell

from b2_automation.cell_evidence import DecisionState, decide_cell

REVIEW_REQUIRED_TEXT = "REVIEW_REQUIRED"
DEFAULT_LOW_CONFIDENCE_THRESHOLD = 0.70


def load_manifest(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _bool_from_spec(value: Any, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    return str(value).strip().lower() in {"1", "true", "yes", "y", "required"}


def _cell_role(spec: Mapping[str, Any]) -> str:
    fid = str(spec.get("field_id", ""))
    role = str(spec.get("cell_role") or spec.get("role") or "target").strip().lower()
    if fid == "evidence_notes" or "note" in role:
        return "notes"
    return role or "target"


def _required_for_spec(spec: Mapping[str, Any], required_field_ids: set[str]) -> bool:
    fid = str(spec.get("field_id", ""))
    if fid in required_field_ids:
        return True
    if "required" in spec:
        return _bool_from_spec(spec.get("required"), default=False)
    return False


def _clear_paragraph_text(paragraph) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""


def _write_cell_preserving_format(cell: _Cell, value: str) -> None:
    """Write a single value without using ``cell.text =``.

    python-docx's cell.text setter clears the cell's paragraph/run XML. This helper
    clears text in existing runs only, preserving paragraph formatting, alignment,
    table cell formatting, and the first run's font properties.
    """
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    _clear_paragraph_text(paragraph)
    paragraph.runs[0].text = str(value)

    # Keep extra paragraphs structurally present but empty so the template shape is
    # not destroyed and stale text cannot survive from prior fills.
    for extra in cell.paragraphs[1:]:
        _clear_paragraph_text(extra)


def _append_notes_preserving_format(cell: _Cell, value: str) -> None:
    """Append multi-line notes while preserving the existing first paragraph style."""
    lines = [line for line in str(value).splitlines()]
    if not lines:
        return
    if not cell.paragraphs:
        base = cell.add_paragraph()
    else:
        base = cell.paragraphs[0]
    target = base if not base.text.strip() else cell.add_paragraph()
    target.style = base.style
    for idx, line in enumerate(lines):
        if idx:
            target.add_run().add_break()
        target.add_run(line)


def _review_text(fid: str, reason: str) -> str:
    return f"{REVIEW_REQUIRED_TEXT}: {fid} {reason}".strip()


def fill_b24_rl1_partial(
    template_path: Path,
    manifest: dict[str, Any],
    field_values: dict[str, str],
    output_path: Path,
    *,
    legacy_only: bool = False,
    field_confidences: Mapping[str, float | None] | None = None,
    required_field_ids: set[str] | None = None,
    low_confidence_threshold: float = DEFAULT_LOW_CONFIDENCE_THRESHOLD,
) -> Path:
    if not legacy_only:
        raise RuntimeError(
            "fill_b24_rl1_partial is legacy-only. Production DOCX writes must use "
            "b2_automation.ooxml_writer.patch_docx_cells with structure guard reporting."
        )
    return fill_b24_rl1_partial_legacy(
        template_path,
        manifest,
        field_values,
        output_path,
        field_confidences=field_confidences,
        required_field_ids=required_field_ids,
        low_confidence_threshold=low_confidence_threshold,
    )


def fill_b24_rl1_partial_legacy(
    template_path: Path,
    manifest: dict[str, Any],
    field_values: dict[str, str],
    output_path: Path,
    *,
    field_confidences: Mapping[str, float | None] | None = None,
    required_field_ids: set[str] | None = None,
    low_confidence_threshold: float = DEFAULT_LOW_CONFIDENCE_THRESHOLD,
) -> Path:
    """
    Apply field_values to cells listed in manifest['cells'].

    Semantics:
    - omitted optional field_id: leave the template cell unchanged.
    - omitted required field_id: write REVIEW_REQUIRED to that cell.
    - present field_id with "": intentionally clear the cell.
    - present field_id with low confidence: write REVIEW_REQUIRED instead of trusting it.
    - evidence/notes cells: append multi-line notes without clearing cell structure.
    """
    confidences = dict(field_confidences or {})
    required_ids = set(required_field_ids or set())

    doc = Document(template_path)
    for spec in manifest.get("cells", []):
        fid = str(spec["field_id"])
        t, r, c = int(spec["table_index"]), int(spec["row"]), int(spec["col"])
        cell = doc.tables[t].rows[r].cells[c]
        role = _cell_role(spec)
        required = _required_for_spec(spec, required_ids)
        value_present = fid in field_values
        value = field_values.get(fid)
        confidence = confidences.get(fid)

        if role == "notes":
            if value_present and value is not None and str(value).strip():
                _append_notes_preserving_format(cell, str(value))
            elif required:
                _write_cell_preserving_format(cell, _review_text(fid, "missing required notes"))
            continue

        if not value_present:
            if required:
                _write_cell_preserving_format(cell, _review_text(fid, "missing required value"))
            continue

        decision = decide_cell(
            value,
            confidence=confidence,
            threshold=low_confidence_threshold,
            required=required,
            conflict_detected=str(value).startswith(REVIEW_REQUIRED_TEXT),
            cell_role="target",
        )
        if decision == DecisionState.BLANK:
            _write_cell_preserving_format(cell, "")
        elif decision in {DecisionState.REVIEW_REQUIRED, DecisionState.MISSING, DecisionState.CONFLICT, DecisionState.LOW_CONFIDENCE}:
            reason = "requires review"
            if value is None or str(value).strip() == "":
                reason = "missing required value"
            elif confidence is not None and confidence < low_confidence_threshold:
                reason = f"low confidence {confidence:.2f}"
            elif str(value).startswith(REVIEW_REQUIRED_TEXT):
                reason = str(value)
            _write_cell_preserving_format(cell, _review_text(fid, reason))
        else:
            _write_cell_preserving_format(cell, str(value))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path.resolve()
