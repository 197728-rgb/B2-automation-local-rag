"""Minimal path: DocuPipe-style JSON -> normalized record -> sample DOCX."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from b2_automation.normalizer import normalize_docupipe_payload


def run_sample_pipeline(input_json: Path, output_docx: Path) -> Path:
    raw = json.loads(input_json.read_text(encoding="utf-8"))
    n = normalize_docupipe_payload(raw)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("B-2 Sample (DocuPipe fixture path)", level=1)
    doc.add_paragraph(f"Schema ID: {n.get('schema_id', '')}")
    doc.add_paragraph(f"Document ID: {n.get('document_id', '')}")
    doc.add_paragraph(f"Status: {n.get('status', '')}")
    doc.add_heading("Normalized fields", level=2)
    doc.add_paragraph(f"Organization: {n.get('organization', '')}")
    if "organization_confidence" in n:
        doc.add_paragraph(f"Organization confidence: {n['organization_confidence']}")
    doc.add_paragraph(f"Audit date: {n.get('audit_date', '')}")
    if "audit_date_confidence" in n:
        doc.add_paragraph(f"Audit date confidence: {n['audit_date_confidence']}")
    doc.add_paragraph(f"Auditor: {n.get('auditor', '')}")
    if "auditor_confidence" in n:
        doc.add_paragraph(f"Auditor confidence: {n['auditor_confidence']}")

    doc.save(output_docx)
    return output_docx.resolve()
