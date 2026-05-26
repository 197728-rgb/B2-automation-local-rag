"""Exhibit B-2 audit validation orchestrator.

Produces JSON + Markdown review artifacts. Does NOT repair table geometry.
Supports ``safe_text_patch_only`` verification via structure guard integration.
"""

from __future__ import annotations

import json
import re
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Mapping

from b2_automation.audit_text_safety import normalize_cell_text, notes_to_dicts
from b2_automation.cell_boundary_validation import validate_cell_boundaries
from b2_automation.docx_structure import extract_docx_structure
from b2_automation.ooxml_writer import build_structure_guard, count_docx_structure
from b2_automation.table_fingerprint import (
    compare_fingerprints,
    fingerprint_docx_tables,
    infer_form_id_from_filename,
    load_expected_fingerprints,
)

ISO_DATE_RE = re.compile(r"\b\d{4}-\d{2}-\d{2}\b")
MYYYY_RE = re.compile(r"(?<!\d/)\b\d{1,2}/\d{4}\b(?!\d)")


@dataclass
class AuditValidationReport:
    docx_path: str
    form_id: str
    validated_at: str
    environment_notes: tuple[str, ...] = (
        "Windows + Office 365 64-bit + Acrobat DC + VM mixed PDF/DOCX workflow assumed",
        "Word merge-cell OOXML treated as fragile; no auto geometry repair",
    )
    safe_text_patch_only: bool = True
    structure_guard: dict[str, Any] | None = None
    layout_integrity_warnings: list[str] = field(default_factory=list)
    normalization_issues: list[dict[str, Any]] = field(default_factory=list)
    fingerprint_comparisons: list[dict[str, Any]] = field(default_factory=list)
    cell_boundary_issues: list[dict[str, Any]] = field(default_factory=list)
    date_format_issues: list[dict[str, Any]] = field(default_factory=list)
    calibration_anomalies: list[dict[str, Any]] = field(default_factory=list)
    flagged_only_inconsistencies: list[dict[str, Any]] = field(default_factory=list)
    typo_fixes_applied: list[dict[str, Any]] = field(default_factory=list)
    unresolved_conflicts: list[dict[str, Any]] = field(default_factory=list)
    visual_export_checks: list[dict[str, Any]] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "docx_path": self.docx_path,
            "form_id": self.form_id,
            "validated_at": self.validated_at,
            "environment_notes": list(self.environment_notes),
            "safe_text_patch_only": self.safe_text_patch_only,
            "structure_guard": self.structure_guard,
            "layout_integrity_warnings": self.layout_integrity_warnings,
            "normalization_issues": self.normalization_issues,
            "fingerprint_comparisons": self.fingerprint_comparisons,
            "cell_boundary_issues": self.cell_boundary_issues,
            "date_format_issues": self.date_format_issues,
            "calibration_anomalies": self.calibration_anomalies,
            "flagged_only_inconsistencies": self.flagged_only_inconsistencies,
            "typo_fixes_applied": self.typo_fixes_applied,
            "unresolved_conflicts": self.unresolved_conflicts,
            "visual_export_checks": self.visual_export_checks,
            "summary": self.summary_dict(),
        }

    def summary_dict(self) -> dict[str, int | bool]:
        fp_drift = sum(1 for x in self.fingerprint_comparisons if x.get("status") == "drift")
        boundary = len(self.cell_boundary_issues)
        dates = len(self.date_format_issues)
        layout = len(self.layout_integrity_warnings)
        flagged = len(self.flagged_only_inconsistencies)
        unresolved = len(self.unresolved_conflicts)
        typos = len(self.typo_fixes_applied)
        return {
            "fingerprint_drift_tables": fp_drift,
            "cell_boundary_issues": boundary,
            "date_format_issues": dates,
            "layout_integrity_warnings": layout,
            "flagged_only_inconsistencies": flagged,
            "typo_fixes_applied": typos,
            "unresolved_conflicts": unresolved,
            "pass": fp_drift == 0 and boundary == 0 and dates == 0 and layout == 0 and unresolved == 0,
        }


def verify_safe_text_patch_only(
    before_path: Path,
    after_path: Path,
) -> dict[str, Any]:
    """Confirm only ``w:t`` nodes changed count; table geometry stable."""
    before = count_docx_structure(before_path)
    after = count_docx_structure(after_path)
    text_delta = after.get("text_nodes", 0) - before.get("text_nodes", 0)
    guard = build_structure_guard(before, after, [], expected_text_node_delta=text_delta)
    guard["safe_text_patch_only"] = guard["pass"]
    guard["mode"] = "safe_text_patch_only"
    return guard


def _scan_document_xml_dates(docx_path: Path) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    for m in ISO_DATE_RE.finditer(xml):
        issues.append(
            {
                "code": "iso_date_in_body",
                "token": m.group(),
                "detail": "YYYY-MM-DD in document body; prefer M/D/YYYY for audit display cells",
            }
        )
    for m in MYYYY_RE.finditer(xml):
        issues.append(
            {
                "code": "month_year_only",
                "token": m.group(),
                "detail": "M/YYYY without day; confirm intended or expand to M/D/YYYY",
            }
        )
    return issues


def _scan_calibration_anomalies(docx_path: Path, form_id: str) -> list[dict[str, Any]]:
    """Flag cal due before cal date in same row when both parse as M/D/Y."""
    from docx import Document

    anomalies: list[dict[str, Any]] = []
    doc = Document(docx_path)
    date_re = re.compile(r"\b(\d{1,2})/(\d{1,2})/(\d{2,4})\b")

    def parse_mdY(s: str):
        m = date_re.search(s)
        if not m:
            return None
        mo, d, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if y < 100:
            y += 2000
        try:
            return datetime(y, mo, d)
        except ValueError:
            return None

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen: set[int] = set()
            dates: list[tuple[int, datetime]] = []
            for ci, cell in enumerate(row.cells):
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                dt = parse_mdY(cell.text)
                if dt:
                    dates.append((ci, dt))
            if len(dates) >= 2:
                dates.sort(key=lambda x: x[1])
                if dates[0][1] > dates[-1][1]:
                    anomalies.append(
                        {
                            "form_location": f"{form_id}/T{ti}R{ri}",
                            "code": "cal_date_order",
                            "detail": "Calibration/due dates in row appear out of chronological order",
                        }
                    )
    return anomalies


def _visual_export_checks(docx_path: Path) -> list[dict[str, Any]]:
    checks: list[dict[str, Any]] = []
    try:
        structure = extract_docx_structure(docx_path)
    except Exception as exc:  # noqa: BLE001
        checks.append({"check": "docx_opens", "pass": False, "detail": str(exc)})
        return checks

    checks.append({"check": "docx_opens", "pass": True, "detail": "python-docx loaded document"})
    checks.append(
        {
            "check": "has_tables",
            "pass": structure.table_count > 0,
            "detail": f"table_count={structure.table_count}",
        }
    )
    checks.append(
        {
            "check": "non_empty_paragraphs",
            "pass": structure.paragraph_count > 0 or structure.table_count > 0,
            "detail": f"paragraph_count={structure.paragraph_count}",
        }
    )
    checks.append(
        {
            "check": "pdf_export_hint",
            "pass": True,
            "detail": "Manual: export PDF in Acrobat/Word and confirm cell alignment visually",
        }
    )
    return checks


def validate_audit_docx(
    docx_path: Path,
    *,
    form_id: str | None = None,
    project_root: Path | None = None,
    reference_template: Path | None = None,
    before_patch_path: Path | None = None,
    typo_fixes_applied: Mapping[str, Any] | None = None,
    flagged_inconsistencies: Mapping[str, Any] | None = None,
    unresolved_conflicts: Mapping[str, Any] | None = None,
    safe_text_patch_only: bool = True,
) -> AuditValidationReport:
    docx_path = Path(docx_path).resolve()
    form_id = form_id or infer_form_id_from_filename(docx_path)
    root = project_root or docx_path.parent

    report = AuditValidationReport(
        docx_path=str(docx_path),
        form_id=form_id,
        validated_at=datetime.now(timezone.utc).replace(microsecond=0).isoformat(),
        safe_text_patch_only=safe_text_patch_only,
    )

    if before_patch_path and before_patch_path.is_file():
        report.structure_guard = verify_safe_text_patch_only(Path(before_patch_path), docx_path)
        if not report.structure_guard.get("pass"):
            report.layout_integrity_warnings.append(
                "Structure guard failed: table/container OOXML counts changed beyond text nodes"
            )

    # Normalization scan on all cell text (flag only)
    try:
        structure = extract_docx_structure(docx_path)
        from docx import Document

        doc = Document(docx_path)
        for ti, table in enumerate(doc.tables):
            seen_cells: set[int] = set()
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    tc_id = id(cell._tc)
                    if tc_id in seen_cells:
                        continue
                    seen_cells.add(tc_id)
                    raw = cell.text
                    if not raw.strip():
                        continue
                    norm = normalize_cell_text(raw)
                    if norm.notes:
                        report.normalization_issues.append(
                            {
                                "form_location": f"{form_id}/T{ti}R{ri}C{ci}",
                                "notes": notes_to_dicts(norm.notes),
                            }
                        )
    except OSError as exc:
        report.layout_integrity_warnings.append(f"Could not scan cells for normalization: {exc}")

    # Fingerprints
    actual_fp = fingerprint_docx_tables(docx_path)
    expected = load_expected_fingerprints(root, form_id)
    if expected is None and reference_template and reference_template.is_file():
        expected = load_expected_fingerprints(reference_template.parent, form_id)
        if expected is None:
            from b2_automation.table_fingerprint import fingerprint_docx_tables as _fp

            ref_tables = _fp(reference_template)
            expected = {
                fp.table_index: {
                    "digest": fp.digest,
                    "row_count": fp.row_count,
                    "header_rows": [list(r) for r in fp.header_rows],
                }
                for fp in ref_tables
            }
    comparisons = compare_fingerprints(actual_fp, expected, form_id=form_id)
    report.fingerprint_comparisons = [c.to_dict() for c in comparisons]
    for c in comparisons:
        if c.status == "drift":
            report.layout_integrity_warnings.append(
                f"Table {c.table_index} header fingerprint drift: {c.detail}"
            )

    report.cell_boundary_issues = [
        i.to_dict() for i in validate_cell_boundaries(docx_path, form_id=form_id)
    ]
    report.date_format_issues = _scan_document_xml_dates(docx_path)
    report.calibration_anomalies = _scan_calibration_anomalies(docx_path, form_id)
    report.visual_export_checks = _visual_export_checks(docx_path)

    if typo_fixes_applied:
        report.typo_fixes_applied = list(typo_fixes_applied.values()) if isinstance(typo_fixes_applied, Mapping) else list(typo_fixes_applied)
    if flagged_inconsistencies:
        report.flagged_only_inconsistencies = (
            list(flagged_inconsistencies.values())
            if isinstance(flagged_inconsistencies, Mapping)
            else list(flagged_inconsistencies)
        )
    if unresolved_conflicts:
        report.unresolved_conflicts = (
            list(unresolved_conflicts.values())
            if isinstance(unresolved_conflicts, Mapping)
            else list(unresolved_conflicts)
        )

    return report


def write_validation_summary(
    report: AuditValidationReport,
    *,
    json_path: Path,
    md_path: Path | None = None,
) -> tuple[Path, Path | None]:
    json_path = Path(json_path)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(report.to_dict(), indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    if md_path is None:
        return json_path, None

    md_path = Path(md_path)
    s = report.summary_dict()
    lines = [
        f"# Audit validation summary — {report.form_id}",
        "",
        f"- **Document:** `{report.docx_path}`",
        f"- **Validated:** {report.validated_at}",
        f"- **Safe text patch only:** {report.safe_text_patch_only}",
        f"- **Overall pass:** {s.get('pass')}",
        "",
        "## Counts",
        "",
        f"| Category | Count |",
        f"|----------|------:|",
        f"| Typo fixes applied | {s.get('typo_fixes_applied', 0)} |",
        f"| Flagged-only inconsistencies | {s.get('flagged_only_inconsistencies', 0)} |",
        f"| Unresolved conflicts | {s.get('unresolved_conflicts', 0)} |",
        f"| Fingerprint drift tables | {s.get('fingerprint_drift_tables', 0)} |",
        f"| Cell boundary issues | {s.get('cell_boundary_issues', 0)} |",
        f"| Date format issues | {s.get('date_format_issues', 0)} |",
        f"| Layout integrity warnings | {s.get('layout_integrity_warnings', 0)} |",
        f"| Calibration anomalies | {len(report.calibration_anomalies)} |",
        "",
    ]

    def _section(title: str, items: list[dict[str, Any]], key: str = "detail") -> None:
        if not items:
            return
        lines.append(f"## {title}")
        lines.append("")
        for item in items[:50]:
            loc = item.get("form_location") or item.get("code") or item.get("check") or "—"
            detail = item.get(key) or item.get("detail") or json.dumps(item, ensure_ascii=False)
            lines.append(f"- **{loc}:** {detail}")
        if len(items) > 50:
            lines.append(f"- … and {len(items) - 50} more (see JSON)")
        lines.append("")

    _section("Typo fixes applied", report.typo_fixes_applied)
    _section("Flagged-only inconsistencies", report.flagged_only_inconsistencies)
    _section("Unresolved conflicts", report.unresolved_conflicts)
    _section("Layout integrity warnings", [{"detail": w} for w in report.layout_integrity_warnings])
    _section("Cell boundary issues", report.cell_boundary_issues)
    _section("Date format issues", report.date_format_issues, key="token")
    _section("Calibration anomalies", report.calibration_anomalies)
    _section("Visual export checks", report.visual_export_checks, key="detail")

    lines.extend(
        [
            "## Environment",
            "",
            "Assumes Windows + Office 365 64-bit + Acrobat DC + VM mixed PDF/DOCX workflow.",
            "Never auto-repair table geometry; suspected corruption is flagged only.",
            "",
        ]
    )
    md_path.write_text("\n".join(lines), encoding="utf-8")
    return json_path, md_path
