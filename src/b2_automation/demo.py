"""Minimal demo: sample DOCX + review JSON (not official AAR templates)."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from b2_automation.paths import resolve_project_root


def run_demo(root: Path | None = None) -> tuple[Path, Path]:
    root = root or resolve_project_root()
    out_dir = root / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_path = out_dir / "sample_b2_form.docx"
    doc = Document()
    doc.add_heading("B-2 Objective Evidence Form", level=1)
    doc.add_paragraph("Organization: ACME Logistics")
    doc.add_paragraph("Audit Date: 2026-05-04")
    doc.add_paragraph("Auditor: Demo Agent")
    doc.add_heading("Evidence Summary", level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item #"
    hdr[1].text = "Finding"
    hdr[2].text = "Status"
    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = "Safety log scanned and verified"
    row1[2].text = "PASS"
    row2 = table.rows[2].cells
    row2[0].text = "2"
    row2[1].text = "Training records on file"
    row2[2].text = "PASS"
    doc.save(docx_path)

    json_path = out_dir / "sample_review.json"
    review = {
        "form": "B-2 Objective Evidence",
        "organization": "ACME Logistics",
        "audit_date": "2026-05-04",
        "items": [
            {"item": 1, "finding": "Safety log scanned and verified", "status": "PASS"},
            {"item": 2, "finding": "Training records on file", "status": "PASS"},
        ],
        "overall_status": "COMPLETE",
    }
    with json_path.open("w", encoding="utf-8") as f:
        json.dump(review, f, indent=2)

    return docx_path, json_path
