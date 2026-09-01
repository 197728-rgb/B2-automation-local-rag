"""Authoritative DOCX structure extraction (python-docx). Mammoth is not used here."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table


@dataclass(frozen=True)
class DocxBlankCell:
    table_index: int
    row_index: int
    column_index: int
    label_text: str
    nearby_header: str
    cell_text: str
    is_blank: bool


@dataclass(frozen=True)
class DocxStructure:
    template_path: str
    paragraph_count: int
    table_count: int
    paragraphs: tuple[str, ...]
    blank_cells: tuple[DocxBlankCell, ...]

    def to_summary_dict(self) -> dict[str, Any]:
        return {
            "template_path": self.template_path,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "blank_cell_count": len(self.blank_cells),
        }


def _cell_text(cell) -> str:
    return re.sub(r"\s+", " ", " ".join(p.text for p in cell.paragraphs)).strip()


def _is_blank(text: str) -> bool:
    if not text:
        return True
    lowered = text.lower()
    return lowered in {"", "n/a", "na", "—", "-", "tbd", "none", "click or tap here to enter text."}


def _row_label(table: Table, row_idx: int, col_idx: int) -> str:
    """Best-effort label from cells to the left or above."""
    parts: list[str] = []
    row = table.rows[row_idx]
    for c in range(min(col_idx, len(row.cells))):
        t = _cell_text(row.cells[c])
        if t and not _is_blank(t):
            parts.append(t)
    if row_idx > 0:
        above = table.rows[row_idx - 1]
        if col_idx < len(above.cells):
            t = _cell_text(above.cells[col_idx])
            if t:
                parts.append(t)
    return " | ".join(parts[:3])


def _table_header(table: Table, max_rows: int = 3) -> str:
    bits: list[str] = []
    for r in range(min(max_rows, len(table.rows))):
        row_bits = [_cell_text(c) for c in table.rows[r].cells[:6]]
        bits.append(" ".join(x for x in row_bits if x))
    return " / ".join(bits)[:240]


def extract_docx_structure(docx_path: Path) -> DocxStructure:
    docx_path = Path(docx_path)
    doc = Document(docx_path)
    paragraphs = tuple(re.sub(r"\s+", " ", p.text).strip() for p in doc.paragraphs if p.text.strip())
    blank_cells: list[DocxBlankCell] = []
    for ti, table in enumerate(doc.tables):
        header = _table_header(table)
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = _cell_text(cell)
                label = _row_label(table, ri, ci)
                if _is_blank(text):
                    blank_cells.append(
                        DocxBlankCell(
                            table_index=ti,
                            row_index=ri,
                            column_index=ci,
                            label_text=label,
                            nearby_header=header,
                            cell_text=text,
                            is_blank=True,
                        )
                    )
    return DocxStructure(
        template_path=str(docx_path),
        paragraph_count=len(doc.paragraphs),
        table_count=len(doc.tables),
        paragraphs=paragraphs,
        blank_cells=tuple(blank_cells),
    )


def extract_mammoth_html(docx_path: Path, max_chars: int = 120_000) -> str:
    """Semantic HTML context only — not authoritative for coordinates."""
    try:
        import mammoth
    except ImportError as exc:
        raise RuntimeError("mammoth is required for autonomous mode: pip install -e '.[autonomous]'") from exc
    with Path(docx_path).open("rb") as f:
        result = mammoth.convert_to_html(f)
    html = result.value or ""
    if len(html) > max_chars:
        return html[:max_chars] + "\n<!-- truncated -->"
    return html


def load_template_manifest_cells(root: Path, template_stem: str) -> list[dict[str, Any]]:
    """Optional manifest from schemas/templates for coordinate hints."""
    manifest_path = root / "schemas" / "templates" / f"{template_stem}.json"
    if not manifest_path.is_file():
        return []
    import json

    data = json.loads(manifest_path.read_text(encoding="utf-8"))
    return list(data.get("cells") or [])
